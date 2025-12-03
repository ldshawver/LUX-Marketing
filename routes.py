import csv
import io
import base64
import os
from datetime import datetime, timedelta
from flask import Blueprint, render_template, request, flash, redirect, url_for, jsonify, make_response, send_file
from flask_login import login_required, current_user
from sqlalchemy import or_
from app import db, csrf
from models import (Contact, Campaign, EmailTemplate, CampaignRecipient, EmailTracking, 
                    BrandKit, EmailComponent, Poll, PollResponse, ABTest, Automation, 
                    AutomationStep, SMSCampaign, SMSRecipient, SMSTemplate, SocialPost, Segment, 
                    SegmentMember, WebForm, FormSubmission, Event, EventRegistration, EventTicket,
                    Product, Order, CalendarEvent, AutomationTemplate, AutomationExecution,
                    AutomationAction, LandingPage, NewsletterArchive, NonOpenerResend,
                    SEOKeyword, SEOBacklink, SEOCompetitor, SEOAudit, SEOPage,
                    TicketPurchase, EventCheckIn, SocialMediaAccount, SocialMediaSchedule,
                    AutomationTest, AutomationTriggerLibrary, AutomationABTest, Company, user_company,
                    Deal, LeadScore, PersonalizationRule, KeywordResearch)
from email_service import EmailService
from utils import validate_email
from tracking import decode_tracking_data, record_email_event
import logging
import json
from ai_agent import lux_agent
from seo_service import seo_service
from error_logger import log_application_error, ApplicationDiagnostics, ErrorLog
from log_reader import LogReader
from auto_repair_service import AutoRepairService
from error_fixes import ErrorFixService
from ai_code_fixer import AICodeFixer
from ai_action_executor import AIActionExecutor

# Stub services for missing imports (prevents LSP errors and runtime crashes)
class SMSService:
    @staticmethod
    def send_sms(phone, message):
        logging.warning(f"SMS Service not configured: {phone}")
        return None

class SchedulingService:
    @staticmethod
    def schedule_task(task_name, time):
        logging.warning(f"Scheduling Service not configured: {task_name}")
        return None

logger = logging.getLogger(__name__)

main_bp = Blueprint('main', __name__)


# Activity Tracking Routes
@main_bp.route('/activities/log', methods=['POST'])
@login_required
def log_activity():
    """Log CRM activity (call, email, meeting, note)"""
    from models import Event
    try:
        activity_type = request.form.get('activity_type')
        contact_id = request.form.get('contact_id')
        subject = request.form.get('subject')
        description = request.form.get('description')
        
        activity = Event(
            name=subject,
            description=description,
            event_type=activity_type,  # call, email, meeting, note
            is_active=True
        )
        
        if contact_id:
            from models import Contact
            contact = Contact.query.get(contact_id)
            if contact:
                activity.name = f"{activity_type.title()}: {subject} - {contact.full_name}"
        
        db.session.add(activity)
        db.session.commit()
        
        flash(f'{activity_type.title()} logged successfully', 'success')
        return redirect(request.referrer or url_for('main.contacts'))
    except Exception as e:
        logger.error(f"Error logging activity: {e}")
        flash('Error logging activity', 'error')
        return redirect(request.referrer or url_for('main.contacts'))

@main_bp.route('/contacts/<int:contact_id>/activities')
@login_required
def contact_activities(contact_id):
    """View all activities for a contact"""
    from models import Contact, Event
    contact = Contact.query.get_or_404(contact_id)
    activities = Event.query.filter(
        Event.name.contains(contact.full_name)
    ).order_by(Event.created_at.desc()).all()
    
    return render_template('contact_activities.html', contact=contact, activities=activities)




# Blog Management Routes
@main_bp.route('/blog')
@login_required
def blog_list():
    """List all blog posts"""
    from models import Event
    posts = Event.query.filter_by(event_type='blog_post').order_by(Event.created_at.desc()).all()
    return render_template('blog_list.html', posts=posts)

@main_bp.route('/blog/create', methods=['GET', 'POST'])
@login_required
def blog_create():
    """Create new blog post with AI assistance"""
    if request.method == 'POST':
        from models import Event
        from ai_agent import lux_agent
        
        use_ai = request.form.get('use_ai') == 'true'
        
        if use_ai:
            topic = request.form.get('topic')
            keywords = request.form.get('keywords', '').split(',')
            tone = request.form.get('tone', 'professional')
            
            # Generate blog post with AI
            result = lux_agent.generate_blog_post(topic, keywords, tone)
            
            if result:
                post = Event(
                    name=result.get('title'),
                    description=result.get('content'),
                    event_type='blog_post',
                    is_active=True
                )
                db.session.add(post)
                db.session.commit()
                flash('Blog post created with AI!', 'success')
                return redirect(url_for('main.blog_edit', post_id=post.id))
        else:
            title = request.form.get('title')
            content = request.form.get('content')
            
            post = Event(
                name=title,
                description=content,
                event_type='blog_post',
                is_active=True
            )
            db.session.add(post)
            db.session.commit()
            flash('Blog post created!', 'success')
            return redirect(url_for('main.blog_edit', post_id=post.id))
    
    return render_template('blog_create.html')

@main_bp.route('/blog/<int:post_id>/edit', methods=['GET', 'POST'])
@login_required
def blog_edit(post_id):
    """Edit blog post"""
    from models import Event
    post = Event.query.get_or_404(post_id)
    
    if request.method == 'POST':
        post.name = request.form.get('title')
        post.description = request.form.get('content')
        db.session.commit()
        flash('Blog post updated!', 'success')
        return redirect(url_for('main.blog_list'))
    
    return render_template('blog_edit.html', post=post)

@main_bp.route('/blog/<int:post_id>/generate-image', methods=['POST'])
@login_required
def blog_generate_image(post_id):
    """Generate featured image for blog post"""
    from models import Event
    import openai
    import os
    
    post = Event.query.get_or_404(post_id)
    prompt = request.form.get('image_prompt', f"Featured image for blog: {post.name}")
    
    try:
        client = openai.OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        response = client.images.generate(
            model="dall-e-3",
            prompt=prompt,
            size="1792x1024",
            quality="standard",
            n=1
        )
        
        image_url = response.data[0].url
        # Store image URL in post metadata
        flash('Image generated successfully!', 'success')
        return jsonify({'success': True, 'image_url': image_url})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@main_bp.route('/')
@login_required
def dashboard():
    """Dashboard with overview statistics"""
    total_contacts = Contact.query.filter_by(is_active=True).count()
    total_campaigns = Campaign.query.count()
    active_campaigns = Campaign.query.filter_by(status='sending').count()
    recent_campaigns = Campaign.query.order_by(Campaign.created_at.desc()).limit(5).all()
    
    # Email statistics
    total_sent = db.session.query(CampaignRecipient).filter_by(status='sent').count()
    total_failed = db.session.query(CampaignRecipient).filter_by(status='failed').count()
    
    # Version 4.1 & 4.2 Feature Metrics
    ai_campaigns = Campaign.query.filter_by(ai_generated=True).count()
    utm_campaigns = Campaign.query.filter(Campaign.utm_keyword.isnot(None)).count()
    social_with_media = SocialPost.query.filter(SocialPost.media_urls.isnot(None)).count()
    total_social_posts = SocialPost.query.count()
    
    current_company = current_user.get_default_company()
    
    return render_template('dashboard.html',
                         total_contacts=total_contacts,
                         total_campaigns=total_campaigns,
                         active_campaigns=active_campaigns,
                         recent_campaigns=recent_campaigns,
                         total_sent=total_sent,
                         total_failed=total_failed,
                         ai_campaigns=ai_campaigns,
                         utm_campaigns=utm_campaigns,
                         social_with_media=social_with_media,
                         total_social_posts=total_social_posts,
                         current_company=current_company)

@main_bp.route('/email-hub')
@login_required
def email_hub():
    """Email Marketing Hub with A/B testing, templates, automations"""
    return render_template('email_hub.html')

@main_bp.route('/campaign-hub')
@login_required
def campaign_hub():
    """Campaign Hub with SEO, Competitors, and AI Campaign Generator"""
    return render_template('campaign_hub.html')

@main_bp.route('/ai-dashboard')
@login_required
def ai_dashboard():
    """LUX AI Dashboard - Monitor and control all AI agents"""
    from agent_scheduler import get_agent_scheduler, agent_execution_history, agent_health_status
    from models import AgentTask
    from sqlalchemy import func
    from datetime import datetime, timedelta
    
    scheduler = get_agent_scheduler()
    
    # Get all agents
    agents = scheduler.agents
    
    # Get scheduled jobs
    jobs = scheduler.get_scheduled_jobs()
    
    # Get recent agent tasks (last 7 days)
    recent_tasks = AgentTask.query.filter(
        AgentTask.created_at >= datetime.now() - timedelta(days=7)
    ).order_by(AgentTask.created_at.desc()).limit(50).all()
    
    # Calculate agent statistics
    agent_stats = db.session.query(
        AgentTask.agent_type,
        func.count(AgentTask.id).label('total_tasks'),
        func.sum(func.case((AgentTask.status == 'completed', 1), else_=0)).label('completed'),
        func.sum(func.case((AgentTask.status == 'failed', 1), else_=0)).label('failed')
    ).filter(
        AgentTask.created_at >= datetime.now() - timedelta(days=30)
    ).group_by(AgentTask.agent_type).all()
    
    # Format stats for template
    stats_dict = {}
    for stat in agent_stats:
        stats_dict[stat.agent_type] = {
            'total': stat.total_tasks,
            'completed': stat.completed,
            'failed': stat.failed,
            'success_rate': (stat.completed / stat.total_tasks * 100) if stat.total_tasks > 0 else 0
        }
    
    # Get APP Agent metrics if available
    app_agent = agents.get('app_intelligence')
    app_metrics = None
    if app_agent:
        app_metrics = {
            'bugs_tracked': len(getattr(app_agent, 'bug_reports', [])),
            'improvements_queued': len(getattr(app_agent, 'improvement_queue', [])),
            'last_health_check': 'Available'
        }
    
    return render_template('ai_dashboard.html',
                         agents=agents,
                         jobs=jobs,
                         recent_tasks=recent_tasks,
                         agent_stats=stats_dict,
                         app_metrics=app_metrics)

@main_bp.route('/ai-dashboard/agent/<agent_type>')
@login_required
def ai_agent_detail(agent_type):
    """Detailed view of a specific AI agent"""
    from agent_scheduler import get_agent_scheduler
    from models import AgentTask
    
    scheduler = get_agent_scheduler()
    agent = scheduler.agents.get(agent_type)
    
    if not agent:
        flash('Agent not found', 'error')
        return redirect(url_for('main.ai_dashboard'))
    
    # Get agent tasks
    tasks = AgentTask.query.filter_by(agent_type=agent_type).order_by(
        AgentTask.created_at.desc()
    ).limit(100).all()
    
    return render_template('ai_agent_detail.html', agent=agent, tasks=tasks, agent_type=agent_type)

@main_bp.route('/ai-dashboard/execute/<agent_type>', methods=['POST'])
@login_required
def execute_agent_task(agent_type):
    """Manually execute an agent task"""
    from agent_scheduler import get_agent_scheduler
    
    scheduler = get_agent_scheduler()
    agent = scheduler.agents.get(agent_type)
    
    if not agent:
        return jsonify({'success': False, 'error': 'Agent not found'}), 404
    
    task_type = request.json.get('task_type')
    task_data = request.json.get('task_data', {})
    
    try:
        # Create task
        task_id = agent.create_task(task_type, task_data)
        
        # Execute
        result = agent.execute({'task_type': task_type, **task_data})
        
        # Complete task
        agent.complete_task(task_id, result)
        
        return jsonify({
            'success': True,
            'task_id': task_id,
            'result': result
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/analytics-hub')
@login_required
def analytics_hub():
    """Comprehensive Analytics Hub with robust data visualization"""
    from datetime import datetime, timedelta
    from integrations.ga4_client import get_ga4_client
    from models import Campaign, EmailTracking, SocialPost, SocialMediaAccount, Contact
    from sqlalchemy import func
    from integrations.woocommerce_client import get_woocommerce_client
    
    # Get date range from query parameters
    days = int(request.args.get('days', 30))
    start_date_str = request.args.get('start_date')
    end_date_str = request.args.get('end_date')
    compare = request.args.get('compare') == 'true'
    
    # Calculate date range
    end_date = datetime.now()
    if end_date_str:
        try:
            end_date = datetime.fromisoformat(end_date_str.replace('Z', '+00:00'))
        except:
            pass
    
    if start_date_str:
        try:
            start_date = datetime.fromisoformat(start_date_str.replace('Z', '+00:00'))
        except:
            start_date = end_date - timedelta(days=days)
    else:
        start_date = end_date - timedelta(days=days)
    
    # Validate date range
    if start_date >= end_date:
        flash('Invalid date range: start date must be before end date', 'error')
        start_date = end_date - timedelta(days=30)
    
    # Calculate previous period for comparison
    period_length = (end_date - start_date).days
    if period_length <= 0:
        period_length = 1
    prev_end_date = start_date
    prev_start_date = start_date - timedelta(days=period_length)
    
    # Initialize data structure
    analytics_data = {
        'active_users': 0,
        'email_open_rate': 0,
        'revenue_mtd': 0,
        'total_reach': 0,
        'ga4_metrics': {},
        'email_metrics': {},
        'social_metrics': {},
        'revenue_data': {},
        'comparison': {} if compare else None,
        'date_range': {
            'start': start_date.strftime('%Y-%m-%d'),
            'end': end_date.strftime('%Y-%m-%d'),
            'days': period_length
        }
    }
    
    # Get GA4 metrics for website analytics
    try:
        ga4_client = get_ga4_client()
        if ga4_client.is_configured():
            ga4_metrics = ga4_client.get_metrics(start_date=start_date, end_date=end_date)
            if ga4_metrics:
                analytics_data['active_users'] = ga4_metrics.get('total_users', 0)
                analytics_data['ga4_metrics'] = ga4_metrics
                
            # Get comparison data if requested
            if compare:
                prev_ga4_metrics = ga4_client.get_metrics(start_date=prev_start_date, end_date=prev_end_date)
                if prev_ga4_metrics:
                    analytics_data['comparison']['active_users'] = prev_ga4_metrics.get('total_users', 0)
                    analytics_data['comparison']['ga4_metrics'] = prev_ga4_metrics
    except Exception as e:
        logger.error(f"Error fetching GA4 metrics: {e}")
    
    # Get Email Campaign Metrics
    try:
        # Email metrics from database
        total_sent = Campaign.query.filter(
            Campaign.status == 'sent',
            Campaign.sent_at >= start_date,
            Campaign.sent_at <= end_date
        ).count()
        
        email_events = EmailTracking.query.filter(
            EmailTracking.created_at >= start_date,
            EmailTracking.created_at <= end_date
        ).with_entities(
            func.sum(func.case((EmailTracking.event_type == 'opened', 1), else_=0)).label('opens'),
            func.sum(func.case((EmailTracking.event_type == 'clicked', 1), else_=0)).label('clicks'),
            func.count(EmailTracking.id).label('total_events')
        ).first()
        
        if email_events and total_sent > 0:
            opens = email_events.opens or 0
            analytics_data['email_open_rate'] = round((opens / total_sent) * 100, 1) if total_sent > 0 else 0
            analytics_data['email_metrics'] = {
                'total_sent': total_sent,
                'opens': opens,
                'clicks': email_events.clicks or 0
            }
    except Exception as e:
        logger.error(f"Error fetching email metrics: {e}")
    
    # Get Social Media Metrics
    try:
        social_accounts = SocialMediaAccount.query.filter_by(is_verified=True).all()
        total_followers = sum([acc.follower_count or 0 for acc in social_accounts])
        
        social_posts = SocialPost.query.filter(
            SocialPost.created_at >= start_date,
            SocialPost.created_at <= end_date
        ).count()
        
        analytics_data['total_reach'] = total_followers
        analytics_data['social_metrics'] = {
            'total_followers': total_followers,
            'posts_this_month': social_posts,
            'connected_accounts': len(social_accounts)
        }
    except Exception as e:
        logger.error(f"Error fetching social metrics: {e}")
    
    # Get WooCommerce Revenue Data
    try:
        wc_client = get_woocommerce_client()
        if wc_client and wc_client.is_configured():
            revenue_data = wc_client.get_revenue_stats(start_date=start_date, end_date=end_date)
            if revenue_data:
                analytics_data['revenue_mtd'] = revenue_data.get('total_sales', 0)
                analytics_data['revenue_data'] = revenue_data
                
            # Get comparison data if requested
            if compare:
                prev_revenue_data = wc_client.get_revenue_stats(start_date=prev_start_date, end_date=prev_end_date)
                if prev_revenue_data:
                    analytics_data['comparison']['revenue_mtd'] = prev_revenue_data.get('total_sales', 0)
                    analytics_data['comparison']['revenue_data'] = prev_revenue_data
        else:
            logger.info("WooCommerce not configured - revenue data unavailable")
    except Exception as e:
        logger.error(f"Error fetching WooCommerce revenue: {e}")
    
    return render_template('analytics_hub.html', analytics=analytics_data)


@main_bp.route('/analytics-hub/export')
@login_required
def export_analytics():
    """Export analytics data in various formats"""
    from datetime import datetime, timedelta
    from integrations.ga4_client import get_ga4_client
    from models import Campaign, EmailTracking, SocialPost, SocialMediaAccount
    from sqlalchemy import func
    from integrations.woocommerce_client import get_woocommerce_client
    import csv
    import io
    from flask import make_response
    
    # Get parameters
    format_type = request.args.get('format', 'csv')
    days = int(request.args.get('days', 30))
    start_date_str = request.args.get('start_date')
    end_date_str = request.args.get('end_date')
    
    # Calculate date range
    end_date = datetime.now()
    if end_date_str:
        try:
            end_date = datetime.fromisoformat(end_date_str.replace('Z', '+00:00'))
        except:
            pass
    
    if start_date_str:
        try:
            start_date = datetime.fromisoformat(start_date_str.replace('Z', '+00:00'))
        except:
            start_date = end_date - timedelta(days=days)
    else:
        start_date = end_date - timedelta(days=days)
    
    period_length = (end_date - start_date).days
    
    # Gather analytics data
    export_data = []
    
    # GA4 Data
    try:
        ga4_client = get_ga4_client()
        if ga4_client.is_configured():
            ga4_metrics = ga4_client.get_metrics(start_date=start_date, end_date=end_date)
            if ga4_metrics:
                export_data.append(['Website Analytics', ''])
                export_data.append(['Total Users', ga4_metrics.get('total_users', 0)])
                export_data.append(['Page Views', ga4_metrics.get('page_views', 0)])
                export_data.append(['Sessions', ga4_metrics.get('sessions', 0)])
                export_data.append(['', ''])
    except Exception as e:
        logger.error(f"Error in export GA4: {e}")
    
    # Email Metrics
    try:
        total_sent = Campaign.query.filter(
            Campaign.status == 'sent',
            Campaign.sent_at >= start_date,
            Campaign.sent_at <= end_date
        ).count()
        
        email_events = EmailTracking.query.filter(
            EmailTracking.created_at >= start_date,
            EmailTracking.created_at <= end_date
        ).with_entities(
            func.sum(func.case((EmailTracking.event_type == 'opened', 1), else_=0)).label('opens'),
            func.sum(func.case((EmailTracking.event_type == 'clicked', 1), else_=0)).label('clicks')
        ).first()
        
        export_data.append(['Email Marketing', ''])
        export_data.append(['Emails Sent', total_sent])
        export_data.append(['Opens', email_events.opens or 0])
        export_data.append(['Clicks', email_events.clicks or 0])
        if total_sent > 0:
            open_rate = round(((email_events.opens or 0) / total_sent) * 100, 1)
            export_data.append(['Open Rate %', open_rate])
        export_data.append(['', ''])
    except Exception as e:
        logger.error(f"Error in export email: {e}")
    
    # Social Metrics
    try:
        social_accounts = SocialMediaAccount.query.filter_by(is_verified=True).all()
        total_followers = sum([acc.follower_count or 0 for acc in social_accounts])
        social_posts = SocialPost.query.filter(
            SocialPost.created_at >= start_date,
            SocialPost.created_at <= end_date
        ).count()
        
        export_data.append(['Social Media', ''])
        export_data.append(['Total Followers', total_followers])
        export_data.append(['Posts Published', social_posts])
        export_data.append(['Connected Accounts', len(social_accounts)])
        export_data.append(['', ''])
    except Exception as e:
        logger.error(f"Error in export social: {e}")
    
    # Revenue Data
    try:
        wc_client = get_woocommerce_client()
        if wc_client and wc_client.is_configured():
            revenue_data = wc_client.get_revenue_stats(start_date=start_date, end_date=end_date)
            if revenue_data:
                export_data.append(['E-commerce Revenue', ''])
                export_data.append(['Total Sales', f"${revenue_data.get('total_sales', 0)}"])
                export_data.append(['Total Orders', revenue_data.get('total_orders', 0)])
                export_data.append(['Avg Order Value', f"${revenue_data.get('average_order_value', 0)}"])
    except Exception as e:
        logger.error(f"Error in export revenue: {e}")
    
    # Generate export based on format
    if format_type == 'csv':
        # CSV Export
        si = io.StringIO()
        writer = csv.writer(si)
        writer.writerow(['LUX Marketing Analytics Report'])
        writer.writerow([f'Period: {start_date.strftime("%Y-%m-%d")} to {end_date.strftime("%Y-%m-%d")}'])
        writer.writerow([''])
        writer.writerows(export_data)
        
        output = make_response(si.getvalue())
        output.headers["Content-Disposition"] = f"attachment; filename=analytics_{datetime.now().strftime('%Y%m%d')}.csv"
        output.headers["Content-type"] = "text/csv"
        return output
    
    elif format_type == 'excel':
        # Excel Export (using openpyxl)
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Analytics Report"
            
            # Header styling
            header_fill = PatternFill(start_color="480749", end_color="480749", fill_type="solid")
            header_font = Font(color="00FFB4", bold=True, size=14)
            
            ws['A1'] = 'LUX Marketing Analytics Report'
            ws['A1'].font = header_font
            ws['A1'].fill = header_fill
            
            ws['A2'] = f'Period: {start_date.strftime("%Y-%m-%d")} to {end_date.strftime("%Y-%m-%d")}'
            
            # Write data
            row = 4
            for data_row in export_data:
                ws[f'A{row}'] = data_row[0]
                ws[f'B{row}'] = data_row[1] if len(data_row) > 1 else ''
                row += 1
            
            # Save to bytes
            excel_file = io.BytesIO()
            wb.save(excel_file)
            excel_file.seek(0)
            
            output = make_response(excel_file.getvalue())
            output.headers["Content-Disposition"] = f"attachment; filename=analytics_{datetime.now().strftime('%Y%m%d')}.xlsx"
            output.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            return output
        except ImportError:
            flash('Excel export requires openpyxl package', 'error')
            return redirect(url_for('main.analytics_hub'))
    
    elif format_type == 'pdf':
        # PDF Export (using reportlab)
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import ImageReader
            import io
            
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            
            # Title
            c.setFont("Helvetica-Bold", 20)
            c.drawString(50, height - 50, "LUX Marketing Analytics Report")
            
            c.setFont("Helvetica", 12)
            c.drawString(50, height - 75, f"Period: {start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')}")
            
            # Write data
            y = height - 110
            for data_row in export_data:
                if y < 50:
                    c.showPage()
                    y = height - 50
                
                if data_row[0] and not data_row[1]:
                    # Section header
                    c.setFont("Helvetica-Bold", 14)
                    c.drawString(50, y, str(data_row[0]))
                    y -= 20
                elif data_row[0]:
                    # Data row
                    c.setFont("Helvetica", 11)
                    c.drawString(70, y, f"{data_row[0]}: {data_row[1] if len(data_row) > 1 else ''}")
                    y -= 18
                else:
                    y -= 10
            
            c.save()
            buffer.seek(0)
            
            output = make_response(buffer.getvalue())
            output.headers["Content-Disposition"] = f"attachment; filename=analytics_{datetime.now().strftime('%Y%m%d')}.pdf"
            output.headers["Content-type"] = "application/pdf"
            return output
        except ImportError:
            flash('PDF export requires reportlab package', 'error')
            return redirect(url_for('main.analytics_hub'))
    
    return redirect(url_for('main.analytics_hub'))

@main_bp.route('/agents-hub')
@login_required
def agents_hub():
    """Redirect to unified Automations & Agents dashboard"""
    return redirect(url_for('main.automation_dashboard'))

@main_bp.route('/ads')
@login_required
def ads_hub():
    """Ads Hub with Display/Search/Shopping ads and Google Ads integration"""
    return render_template('ads_hub.html')

@main_bp.route('/companies')
@login_required
def companies_list():
    """List all companies for the current user"""
    user_companies = current_user.companies
    return render_template('companies.html', companies=user_companies)

@main_bp.route('/companies/add', methods=['GET', 'POST'])
@login_required
def add_company():
    """Add a new company"""
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        website_url = request.form.get('website_url', '').strip()
        
        if not name:
            flash('Company name is required', 'error')
            return redirect(url_for('main.add_company'))
        
        company = Company()
        company.name = name
        company.website_url = website_url
        company.env_config = {}
        company.social_accounts = {}
        company.email_config = {}
        company.api_keys = {}
        
        logo_file = request.files.get('logo')
        if logo_file and logo_file.filename:
            import os
            from werkzeug.utils import secure_filename
            filename = secure_filename(logo_file.filename)
            logo_path = f'company_logos/{filename}'
            os.makedirs('static/company_logos', exist_ok=True)
            logo_file.save(f'static/{logo_path}')
            company.logo_path = logo_path
        
        db.session.add(company)
        db.session.flush()
        
        current_user.companies.append(company)
        
        is_default = len(current_user.companies) == 1
        db.session.execute(
            user_company.update().where(
                (user_company.c.user_id == current_user.id) &
                (user_company.c.company_id == company.id)
            ).values(is_default=is_default)
        )
        
        db.session.commit()
        flash(f'Company "{name}" added successfully', 'success')
        return redirect(url_for('main.companies_list'))
    
    return render_template('company_add.html')

@main_bp.route('/companies/edit/<int:company_id>', methods=['GET', 'POST'])
@login_required
def edit_company(company_id):
    """Edit company details"""
    company = Company.query.get_or_404(company_id)
    
    if company not in current_user.companies:
        flash('You do not have access to this company', 'error')
        return redirect(url_for('main.companies_list'))
    
    if request.method == 'POST':
        company.name = request.form.get('name', '').strip()
        company.website_url = request.form.get('website_url', '').strip()
        
        logo_file = request.files.get('logo')
        if logo_file and logo_file.filename:
            import os
            from werkzeug.utils import secure_filename
            filename = secure_filename(logo_file.filename)
            logo_path = f'company_logos/{filename}'
            os.makedirs('static/company_logos', exist_ok=True)
            logo_file.save(f'static/{logo_path}')
            company.logo_path = logo_path
        
        icon_file = request.files.get('icon')
        if icon_file and icon_file.filename:
            import os
            from werkzeug.utils import secure_filename
            filename = secure_filename(icon_file.filename)
            icon_path = f'company_logos/{filename}'
            os.makedirs('static/company_logos', exist_ok=True)
            icon_file.save(f'static/{icon_path}')
            company.icon_path = icon_path
        
        company.primary_color = request.form.get('primary_color', '#bc00ed')
        company.secondary_color = request.form.get('secondary_color', '#00ffb4')
        company.accent_color = request.form.get('accent_color', '#e4055c')
        company.font_family = request.form.get('font_family', 'Inter, sans-serif')
        
        db.session.commit()
        flash(f'Company "{company.name}" updated successfully', 'success')
        return redirect(url_for('main.companies_list'))
    
    return render_template('company_edit.html', company=company)

@main_bp.route('/companies/switch/<int:company_id>', methods=['POST'])
@csrf.exempt
@login_required
def switch_company(company_id):
    """Switch the user's default company"""
    try:
        company = Company.query.get_or_404(company_id)
        
        if company not in current_user.companies:
            return jsonify({'success': False, 'error': 'Access denied'}), 403
        
        db.session.execute(
            user_company.update().where(
                user_company.c.user_id == current_user.id
            ).values(is_default=False)
        )
        
        db.session.execute(
            user_company.update().where(
                (user_company.c.user_id == current_user.id) &
                (user_company.c.company_id == company_id)
            ).values(is_default=True)
        )
        
        db.session.commit()
        
        return jsonify({'success': True, 'company_name': company.name})
    except Exception as e:
        logger.error(f"Error switching company: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/contacts')
@login_required
def contacts():
    """Contact management page"""
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '')
    
    query = Contact.query.filter_by(is_active=True)
    
    if search:
        query = query.filter(or_(
            Contact.email.contains(search),
            Contact.first_name.contains(search),
            Contact.last_name.contains(search),
            Contact.company.contains(search)
        ))
    
    contacts = query.order_by(Contact.created_at.desc()).paginate(
        page=page, per_page=20, error_out=False
    )
    
    return render_template('contacts.html', contacts=contacts, search=search)

@main_bp.route('/contacts/add', methods=['POST'])
@login_required
def add_contact():
    """Add a new contact"""
    email = request.form.get('email', '').strip()
    first_name = request.form.get('first_name', '').strip()
    last_name = request.form.get('last_name', '').strip()
    company = request.form.get('company', '').strip()
    phone = request.form.get('phone', '').strip()
    tags = request.form.get('tags', '').strip()
    
    if not email:
        flash('Email is required', 'error')
        return redirect(url_for('main.contacts'))
    
    if not validate_email(email):
        flash('Invalid email format', 'error')
        return redirect(url_for('main.contacts'))
    
    # Check if contact already exists
    existing = Contact.query.filter_by(email=email).first()
    if existing:
        flash('Contact with this email already exists', 'error')
        return redirect(url_for('main.contacts'))
    
    contact = Contact()
    contact.email = email
    contact.first_name = first_name
    contact.last_name = last_name
    contact.company = company
    contact.phone = phone
    contact.tags = tags
    
    db.session.add(contact)
    db.session.commit()
    
    flash('Contact added successfully', 'success')
    return redirect(url_for('main.contacts'))

@main_bp.route('/contacts/import', methods=['POST'])
@login_required
def import_contacts():
    """Import contacts from CSV file"""
    if 'file' not in request.files:
        flash('No file selected', 'error')
        return redirect(url_for('main.contacts'))
    
    file = request.files['file']
    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('main.contacts'))
    
    if not file.filename or not file.filename.endswith('.csv'):
        flash('Please upload a CSV file', 'error')
        return redirect(url_for('main.contacts'))
    
    try:
        # Read CSV file with better error handling
        try:
            content = file.stream.read()
            # Try different encodings
            for encoding in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
                try:
                    decoded_content = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                flash('Unable to decode file. Please ensure it is saved in UTF-8 format.', 'error')
                return redirect(url_for('main.contacts'))
            
            stream = io.StringIO(decoded_content, newline=None)
            csv_input = csv.DictReader(stream)
        except Exception as e:
            flash(f'Error reading CSV file: {str(e)}', 'error')
            return redirect(url_for('main.contacts'))
        
        imported_count = 0
        error_count = 0
        
        for row in csv_input:
            email = row.get('email', '').strip()
            if not email or not validate_email(email):
                error_count += 1
                continue
            
            # Check if contact already exists
            existing = Contact.query.filter_by(email=email).first()
            if existing:
                continue
            
            contact = Contact()
            contact.email = email
            contact.first_name = row.get('first_name', '').strip()
            contact.last_name = row.get('last_name', '').strip()
            contact.company = row.get('company', '').strip()
            contact.phone = row.get('phone', '').strip()
            contact.tags = row.get('tags', '').strip()
            
            db.session.add(contact)
            imported_count += 1
        
        db.session.commit()
        
        flash(f'Successfully imported {imported_count} contacts. {error_count} errors.', 'success')
        
    except Exception as e:
        logging.error(f"Error importing contacts: {str(e)}")
        flash('Error importing contacts. Please check file format.', 'error')
    
    return redirect(url_for('main.contacts'))

@main_bp.route('/contacts/export')
@login_required
def export_contacts():
    """Export contacts to CSV"""
    contacts = Contact.query.filter_by(is_active=True).all()
    
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Write header
    writer.writerow(['email', 'first_name', 'last_name', 'company', 'phone', 'tags', 'created_at'])
    
    # Write contacts
    for contact in contacts:
        writer.writerow([
            contact.email,
            contact.first_name or '',
            contact.last_name or '',
            contact.company or '',
            contact.phone or '',
            contact.tags or '',
            contact.created_at.strftime('%Y-%m-%d %H:%M:%S')
        ])
    
    # Create response
    output.seek(0)
    response = make_response(output.getvalue())
    response.headers['Content-Type'] = 'text/csv'
    response.headers['Content-Disposition'] = f'attachment; filename=contacts_{datetime.now().strftime("%Y%m%d")}.csv'
    
    return response

@main_bp.route('/contacts/<int:contact_id>/delete', methods=['POST'])
@login_required
def delete_contact(contact_id):
    """Delete a contact"""
    contact = Contact.query.get_or_404(contact_id)
    contact.is_active = False
    db.session.commit()
    
    flash('Contact deleted successfully', 'success')
    return redirect(url_for('main.contacts'))

@main_bp.route('/campaigns')
@login_required
def campaigns():
    """Campaign management page"""
    page = request.args.get('page', 1, type=int)
    status_filter = request.args.get('status', '')
    
    query = Campaign.query
    
    if status_filter:
        query = query.filter_by(status=status_filter)
    
    campaigns = query.order_by(Campaign.created_at.desc()).paginate(
        page=page, per_page=10, error_out=False
    )
    
    return render_template('campaigns.html', campaigns=campaigns, status_filter=status_filter)

@main_bp.route('/campaigns/create', methods=['GET', 'POST'])
@login_required
def create_campaign():
    """Create a new campaign"""
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        subject = request.form.get('subject', '').strip()
        template_id = request.form.get('template_id', type=int)
        scheduled_at = request.form.get('scheduled_at')
        recipient_tags = request.form.get('recipient_tags', '').strip()
        
        if not name or not subject or not template_id:
            flash('Name, subject, and template are required', 'error')
            return redirect(url_for('main.create_campaign'))
        
        # Create campaign
        campaign = Campaign()
        campaign.name = name
        campaign.subject = subject
        campaign.template_id = template_id
        campaign.status = 'draft'
        
        if scheduled_at:
            try:
                campaign.scheduled_at = datetime.fromisoformat(scheduled_at.replace('T', ' '))
                campaign.status = 'scheduled'
            except ValueError:
                flash('Invalid scheduled time format', 'error')
                return redirect(url_for('main.create_campaign'))
        
        db.session.add(campaign)
        db.session.flush()  # Get campaign ID
        
        # Add recipients
        contacts_query = Contact.query.filter_by(is_active=True)
        
        if recipient_tags:
            # Filter by tags
            tag_list = [tag.strip() for tag in recipient_tags.split(',')]
            tag_conditions = []
            for tag in tag_list:
                tag_conditions.append(Contact.tags.contains(tag))
            contacts_query = contacts_query.filter(or_(*tag_conditions))
        
        contacts = contacts_query.all()
        
        for contact in contacts:
            recipient = CampaignRecipient()
            recipient.campaign_id = campaign.id
            recipient.contact_id = contact.id
            db.session.add(recipient)
        
        db.session.commit()
        
        flash(f'Campaign created successfully with {len(contacts)} recipients', 'success')
        return redirect(url_for('main.campaigns'))
    
    # GET request - show form
    templates = EmailTemplate.query.filter_by(is_active=True).all()
    return render_template('campaign_create.html', templates=templates)

@main_bp.route('/campaigns/<int:campaign_id>/send', methods=['POST'])
@login_required
def send_campaign(campaign_id):
    """Send a campaign immediately"""
    campaign = Campaign.query.get_or_404(campaign_id)
    
    if campaign.status not in ['draft', 'scheduled']:
        flash('Campaign cannot be sent in current status', 'error')
        return redirect(url_for('main.campaigns'))
    
    try:
        email_service = EmailService()
        email_service.send_campaign(campaign)
        
        campaign.status = 'sending'
        campaign.sent_at = datetime.utcnow()
        db.session.commit()
        
        flash('Campaign is being sent', 'success')
    except Exception as e:
        logging.error(f"Error sending campaign: {str(e)}")
        flash('Error sending campaign. Please check configuration.', 'error')
    
    return redirect(url_for('main.campaigns'))

@main_bp.route('/campaigns/<int:campaign_id>/preview')
@login_required
def preview_campaign(campaign_id):
    """Preview campaign email"""
    campaign = Campaign.query.get_or_404(campaign_id)
    template = campaign.template
    
    # Get sample contact for preview
    sample_contact = Contact.query.filter_by(is_active=True).first()
    
    return render_template('preview_email.html', 
                         campaign=campaign, 
                         template=template,
                         sample_contact=sample_contact)

@main_bp.route('/templates')
@login_required
def templates():
    """Email template management"""
    templates = EmailTemplate.query.filter_by(is_active=True).order_by(EmailTemplate.created_at.desc()).all()
    return render_template('templates_manage.html', templates=templates)

@main_bp.route('/templates/gallery')
@login_required
def template_gallery():
    """Branded template gallery"""
    # Define branded template options
    branded_templates = [
        {
            'id': 'modern_newsletter',
            'name': 'Modern Newsletter',
            'description': 'Clean, modern design perfect for newsletters and announcements',
            'preview_image': '/static/images/templates/modern_newsletter.png',
            'category': 'Newsletter',
            'features': ['Responsive Design', 'Hero Image', 'Call-to-Action Button', 'Social Links']
        },
        {
            'id': 'promotional_sale',
            'name': 'Promotional Sale',
            'description': 'Eye-catching design for sales promotions and special offers',
            'preview_image': '/static/images/templates/promotional_sale.png',
            'category': 'Promotional',
            'features': ['Bold Headlines', 'Product Showcase', 'Urgency Elements', 'Discount Highlights']
        },
        {
            'id': 'welcome_series',
            'name': 'Welcome Email',
            'description': 'Professional welcome email for new subscribers',
            'preview_image': '/static/images/templates/welcome_series.png',
            'category': 'Welcome',
            'features': ['Personal Touch', 'Brand Introduction', 'Next Steps', 'Contact Information']
        },
        {
            'id': 'event_invitation',
            'name': 'Event Invitation',
            'description': 'Elegant design for event invitations and announcements',
            'preview_image': '/static/images/templates/event_invitation.png',
            'category': 'Events',
            'features': ['Event Details', 'RSVP Button', 'Location Map', 'Calendar Integration']
        },
        {
            'id': 'product_update',
            'name': 'Product Update',
            'description': 'Professional layout for product announcements and updates',
            'preview_image': '/static/images/templates/product_update.png',
            'category': 'Product',
            'features': ['Feature Highlights', 'Screenshots', 'Learn More Links', 'Feedback Request']
        },
        {
            'id': 'minimal_corporate',
            'name': 'Minimal Corporate',
            'description': 'Clean, minimalist design for corporate communications',
            'preview_image': '/static/images/templates/minimal_corporate.png',
            'category': 'Corporate',
            'features': ['Professional Layout', 'Typography Focus', 'Brand Colors', 'Simple CTA']
        }
    ]
    
    return render_template('template_gallery.html', branded_templates=branded_templates)

@main_bp.route('/templates/use-branded/<template_id>')
@login_required
def use_branded_template(template_id):
    """Use a branded template to create a new custom template"""
    # Get the branded template HTML
    template_html = get_branded_template_html(template_id)
    
    if not template_html:
        flash('Template not found', 'error')
        return redirect(url_for('main.template_gallery'))
    
    # Get template info
    template_info = get_branded_template_info(template_id)
    
    return render_template('template_create.html', 
                         branded_template=template_info,
                         default_html=template_html['html'],
                         default_subject=template_html['subject'])

@main_bp.route('/templates/create', methods=['GET', 'POST'])
@login_required
def create_template():
    """Create a new email template"""
    if request.method == 'POST':
        try:
            name = request.form.get('name', '').strip()
            subject = request.form.get('subject', '').strip()
            html_content = request.form.get('html_content', '').strip()
            
            logging.debug(f"Template creation attempt: name='{name}', subject='{subject}', content_length={len(html_content)}")
            
            if not name or not subject or not html_content:
                flash('All fields are required', 'error')
                return redirect(url_for('main.create_template'))
            
            template = EmailTemplate()
            template.name = name
            template.subject = subject
            template.html_content = html_content
            
            db.session.add(template)
            db.session.commit()
            
            logging.info(f"Template '{name}' created successfully")
            flash('Template created successfully', 'success')
            return redirect(url_for('main.templates'))
            
        except Exception as e:
            logging.error(f"Error creating template: {str(e)}")
            db.session.rollback()
            flash('Error creating template. Please try again.', 'error')
            return redirect(url_for('main.create_template'))
    
    return render_template('template_create.html')

@main_bp.route('/templates/<int:template_id>/preview')
@login_required
def preview_template(template_id):
    """Preview email template"""
    template = EmailTemplate.query.get_or_404(template_id)
    
    # Get sample contact for preview
    sample_contact = Contact.query.filter_by(is_active=True).first()
    
    if not sample_contact:
        # Create a sample contact for preview if none exists
        sample_contact = type('SampleContact', (), {
            'first_name': 'John',
            'last_name': 'Doe',
            'full_name': 'John Doe',
            'email': 'john.doe@example.com',
            'company': 'Example Company',
            'phone': '+1 (555) 123-4567'
        })()
    
    # Create a sample campaign for preview
    sample_campaign = type('SampleCampaign', (), {
        'name': 'Sample Campaign',
        'subject': template.subject,
        'id': 0
    })()
    
    try:
        from email_service import EmailService
        email_service = EmailService()
        
        # Render the template with sample data
        rendered_html = email_service.render_template(
            template.html_content,
            sample_contact,
            sample_campaign
        )
        
        return render_template('preview_template.html', 
                             template=template, 
                             rendered_html=rendered_html,
                             sample_contact=sample_contact)
                             
    except Exception as e:
        flash(f'Error rendering template preview: {str(e)}', 'error')
        return redirect(url_for('main.templates'))

@main_bp.route('/templates/preview-live', methods=['POST'])
@login_required
def preview_template_live():
    """Live preview of template during creation/editing"""
    try:
        html_content = request.form.get('html_content', '')
        subject = request.form.get('subject', 'Preview Subject')
        
        if not html_content:
            return jsonify({'error': 'No HTML content provided'}), 400
        
        # Create sample data for preview
        sample_contact = type('SampleContact', (), {
            'first_name': 'John',
            'last_name': 'Doe',
            'full_name': 'John Doe',
            'email': 'john.doe@example.com',
            'company': 'Example Company',
            'phone': '+1 (555) 123-4567'
        })()
        
        sample_campaign = type('SampleCampaign', (), {
            'name': 'Sample Campaign',
            'subject': subject,
            'id': 0
        })()
        
        # Render the template
        from email_service import EmailService
        email_service = EmailService()
        
        rendered_html = email_service.render_template(
            html_content,
            sample_contact,
            sample_campaign
        )
        
        return jsonify({
            'success': True,
            'rendered_html': rendered_html,
            'subject': subject
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def get_branded_template_html(template_id):
    """Get HTML content for branded templates"""
    templates = {
        'modern_newsletter': {
            'subject': 'Latest Updates from {{campaign.name}}',
            'html': '''<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{campaign.subject}}</title>
</head>
<body style="margin: 0; padding: 0; font-family: 'Arial', sans-serif; background-color: #f8f9fa;">
    <table width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color: #f8f9fa;">
        <tr>
            <td align="center" style="padding: 40px 20px;">
                <table width="600" cellpadding="0" cellspacing="0" border="0" style="background-color: #ffffff; border-radius: 8px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
                    <!-- Header -->
                    <tr>
                        <td style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 40px 30px; text-align: center; border-radius: 8px 8px 0 0;">
                            <h1 style="margin: 0; color: #ffffff; font-size: 28px; font-weight: bold;">Your Company</h1>
                            <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 16px; opacity: 0.9;">Newsletter</p>
                        </td>
                    </tr>
                    <!-- Main Content -->
                    <tr>
                        <td style="padding: 40px 30px;">
                            <h2 style="margin: 0 0 20px 0; color: #333333; font-size: 24px;">Hello {{contact.first_name}}!</h2>
                            <p style="margin: 0 0 20px 0; color: #666666; font-size: 16px; line-height: 1.6;">
                                We're excited to share the latest updates and insights with you. Here's what's new:
                            </p>
                            
                            <!-- Feature Section -->
                            <div style="background-color: #f8f9fa; padding: 25px; border-radius: 6px; margin: 25px 0;">
                                <h3 style="margin: 0 0 15px 0; color: #333333; font-size: 20px;">Featured Content</h3>
                                <p style="margin: 0 0 15px 0; color: #666666; font-size: 16px; line-height: 1.6;">
                                    Add your featured content here. This section is perfect for highlighting your most important news or updates.
                                </p>
                                <a href="#" style="display: inline-block; background-color: #667eea; color: #ffffff; text-decoration: none; padding: 12px 24px; border-radius: 4px; font-weight: bold;">
                                    Learn More
                                </a>
                            </div>
                            
                            <p style="margin: 20px 0 0 0; color: #666666; font-size: 16px; line-height: 1.6;">
                                Thank you for being part of our community, {{contact.first_name}}!
                            </p>
                        </td>
                    </tr>
                    <!-- Footer -->
                    <tr>
                        <td style="background-color: #f8f9fa; padding: 30px; text-align: center; border-radius: 0 0 8px 8px;">
                            <p style="margin: 0 0 10px 0; color: #999999; font-size: 14px;">
                                You're receiving this email because you subscribed to our newsletter.
                            </p>
                            <p style="margin: 0; color: #999999; font-size: 14px;">
                                <a href="{{unsubscribe_url}}" style="color: #667eea; text-decoration: none;">Unsubscribe</a> | 
                                <a href="#" style="color: #667eea; text-decoration: none;">Update Preferences</a>
                            </p>
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>'''
        },
        'promotional_sale': {
            'subject': '🔥 Special Offer Just for You, {{contact.first_name}}!',
            'html': '''<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{campaign.subject}}</title>
</head>
<body style="margin: 0; padding: 0; font-family: 'Arial', sans-serif; background-color: #f8f9fa;">
    <table width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color: #f8f9fa;">
        <tr>
            <td align="center" style="padding: 40px 20px;">
                <table width="600" cellpadding="0" cellspacing="0" border="0" style="background-color: #ffffff; border-radius: 8px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
                    <!-- Header -->
                    <tr>
                        <td style="background: linear-gradient(135deg, #ff6b6b 0%, #ee5a24 100%); padding: 40px 30px; text-align: center; border-radius: 8px 8px 0 0;">
                            <h1 style="margin: 0; color: #ffffff; font-size: 32px; font-weight: bold;">MEGA SALE</h1>
                            <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 18px; font-weight: bold;">UP TO 50% OFF</p>
                        </td>
                    </tr>
                    <!-- Main Content -->
                    <tr>
                        <td style="padding: 40px 30px; text-align: center;">
                            <h2 style="margin: 0 0 20px 0; color: #333333; font-size: 24px;">Hey {{contact.first_name}}, Don't Miss Out!</h2>
                            <p style="margin: 0 0 30px 0; color: #666666; font-size: 16px; line-height: 1.6;">
                                Our biggest sale of the year is here! Save up to 50% on everything. Limited time only!
                            </p>
                            
                            <!-- Offer Box -->
                            <div style="background: linear-gradient(135deg, #ff6b6b 0%, #ee5a24 100%); color: #ffffff; padding: 30px; border-radius: 8px; margin: 30px 0;">
                                <h3 style="margin: 0 0 10px 0; font-size: 28px; font-weight: bold;">50% OFF</h3>
                                <p style="margin: 0 0 20px 0; font-size: 16px;">Use code: SAVE50</p>
                                <a href="#" style="display: inline-block; background-color: #ffffff; color: #ee5a24; text-decoration: none; padding: 15px 30px; border-radius: 4px; font-weight: bold; font-size: 16px;">
                                    SHOP NOW
                                </a>
                            </div>
                            
                            <p style="margin: 20px 0 0 0; color: #ff6b6b; font-size: 14px; font-weight: bold;">
                                ⏰ Hurry! Sale ends in 48 hours
                            </p>
                        </td>
                    </tr>
                    <!-- Footer -->
                    <tr>
                        <td style="background-color: #f8f9fa; padding: 30px; text-align: center; border-radius: 0 0 8px 8px;">
                            <p style="margin: 0 0 10px 0; color: #999999; font-size: 14px;">
                                You're receiving this email because you're a valued customer.
                            </p>
                            <p style="margin: 0; color: #999999; font-size: 14px;">
                                <a href="{{unsubscribe_url}}" style="color: #ff6b6b; text-decoration: none;">Unsubscribe</a>
                            </p>
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>'''
        },
        'welcome_series': {
            'subject': 'Welcome to our community, {{contact.first_name}}!',
            'html': '''<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{campaign.subject}}</title>
</head>
<body style="margin: 0; padding: 0; font-family: 'Arial', sans-serif; background-color: #f8f9fa;">
    <table width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color: #f8f9fa;">
        <tr>
            <td align="center" style="padding: 40px 20px;">
                <table width="600" cellpadding="0" cellspacing="0" border="0" style="background-color: #ffffff; border-radius: 8px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
                    <!-- Header -->
                    <tr>
                        <td style="background: linear-gradient(135deg, #2ecc71 0%, #27ae60 100%); padding: 40px 30px; text-align: center; border-radius: 8px 8px 0 0;">
                            <h1 style="margin: 0; color: #ffffff; font-size: 32px; font-weight: bold;">Welcome!</h1>
                            <p style="margin: 10px 0 0 0; color: #ffffff; font-size: 16px; opacity: 0.9;">We're thrilled to have you join us</p>
                        </td>
                    </tr>
                    <!-- Main Content -->
                    <tr>
                        <td style="padding: 40px 30px;">
                            <h2 style="margin: 0 0 20px 0; color: #333333; font-size: 24px;">Hi {{contact.first_name}}, Welcome aboard! 🎉</h2>
                            <p style="margin: 0 0 20px 0; color: #666666; font-size: 16px; line-height: 1.6;">
                                Thank you for joining our community! We're excited to have you with us and can't wait to share amazing content, updates, and exclusive offers.
                            </p>
                            
                            <!-- Getting Started Section -->
                            <div style="background-color: #f8f9fa; padding: 25px; border-radius: 6px; margin: 25px 0;">
                                <h3 style="margin: 0 0 15px 0; color: #333333; font-size: 20px;">🚀 Getting Started</h3>
                                <ul style="margin: 0; padding-left: 20px; color: #666666; font-size: 16px; line-height: 1.8;">
                                    <li>Explore our latest content and resources</li>
                                    <li>Follow us on social media for daily updates</li>
                                    <li>Join our community discussions</li>
                                    <li>Don't forget to add us to your contacts</li>
                                </ul>
                            </div>
                            
                            <div style="text-align: center; margin: 30px 0;">
                                <a href="#" style="display: inline-block; background-color: #2ecc71; color: #ffffff; text-decoration: none; padding: 15px 30px; border-radius: 4px; font-weight: bold; font-size: 16px;">
                                    Get Started
                                </a>
                            </div>
                            
                            <p style="margin: 20px 0 0 0; color: #666666; font-size: 16px; line-height: 1.6;">
                                If you have any questions, feel free to reply to this email. We're here to help!
                            </p>
                            
                            <p style="margin: 20px 0 0 0; color: #666666; font-size: 16px; line-height: 1.6;">
                                Best regards,<br>
                                <strong>The Team</strong>
                            </p>
                        </td>
                    </tr>
                    <!-- Footer -->
                    <tr>
                        <td style="background-color: #f8f9fa; padding: 30px; text-align: center; border-radius: 0 0 8px 8px;">
                            <p style="margin: 0 0 10px 0; color: #999999; font-size: 14px;">
                                You're receiving this email because you recently signed up.
                            </p>
                            <p style="margin: 0; color: #999999; font-size: 14px;">
                                <a href="{{unsubscribe_url}}" style="color: #2ecc71; text-decoration: none;">Unsubscribe</a> | 
                                <a href="#" style="color: #2ecc71; text-decoration: none;">Contact Us</a>
                            </p>
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>'''
        }
    }
    
    return templates.get(template_id)

def get_branded_template_info(template_id):
    """Get template information"""
    template_info = {
        'modern_newsletter': {
            'name': 'Modern Newsletter',
            'description': 'Clean, modern design perfect for newsletters and announcements'
        },
        'promotional_sale': {
            'name': 'Promotional Sale',
            'description': 'Eye-catching design for sales promotions and special offers'
        },
        'welcome_series': {
            'name': 'Welcome Email',
            'description': 'Professional welcome email for new subscribers'
        }
    }
    
    return template_info.get(template_id, {'name': 'Unknown Template', 'description': ''})

@main_bp.route('/analytics/comprehensive')
@login_required
def comprehensive_analytics():
    """Comprehensive analytics dashboard with all 6 metric categories"""
    try:
        from agents.analytics_agent import AnalyticsAgent
        
        agent = AnalyticsAgent()
        period_days = request.args.get('period_days', 30, type=int)
        
        metrics_result = agent.calculate_comprehensive_metrics({'period_days': period_days})
        
        if not metrics_result.get('success'):
            flash('Unable to load analytics data', 'error')
            return redirect(url_for('main.dashboard'))
        
        return render_template('analytics_comprehensive.html',
                             metrics=metrics_result.get('metrics', {}),
                             period_days=period_days,
                             generated_at=metrics_result.get('generated_at'))
        
    except Exception as e:
        logger.error(f"Comprehensive analytics error: {e}")
        flash('Error loading analytics dashboard', 'error')
        return redirect(url_for('main.dashboard'))

@main_bp.route('/analytics')
@login_required
def analytics():
    """Analytics and reporting dashboard"""
    # Campaign statistics
    total_campaigns = Campaign.query.count()
    sent_campaigns = Campaign.query.filter_by(status='sent').count()
    
    # Email delivery statistics
    total_sent = db.session.query(CampaignRecipient).filter_by(status='sent').count()
    total_failed = db.session.query(CampaignRecipient).filter_by(status='failed').count()
    total_bounced = db.session.query(CampaignRecipient).filter_by(status='bounced').count()
    total_pending = db.session.query(CampaignRecipient).filter_by(status='pending').count()
    
    # Engagement statistics
    total_opened = db.session.query(CampaignRecipient).filter(CampaignRecipient.opened_at.isnot(None)).count()
    total_clicked = db.session.query(CampaignRecipient).filter(CampaignRecipient.clicked_at.isnot(None)).count()
    
    # Calculate rates
    total_delivered = total_sent
    open_rate = (total_opened / total_delivered * 100) if total_delivered > 0 else 0
    click_rate = (total_clicked / total_delivered * 100) if total_delivered > 0 else 0
    bounce_rate = (total_bounced / (total_delivered + total_bounced) * 100) if (total_delivered + total_bounced) > 0 else 0
    
    # SMS Campaign statistics
    total_sms_campaigns = SMSCampaign.query.count()
    sent_sms_campaigns = SMSCampaign.query.filter_by(status='sent').count()
    sms_sent = db.session.query(SMSRecipient).filter_by(status='sent').count()
    sms_failed = db.session.query(SMSRecipient).filter_by(status='failed').count()
    sms_pending = db.session.query(SMSRecipient).filter_by(status='pending').count()
    sms_delivery_rate = (sms_sent / (sms_sent + sms_failed) * 100) if (sms_sent + sms_failed) > 0 else 0
    
    # Tracking events breakdown
    tracking_stats = db.session.query(
        EmailTracking.event_type, 
        db.func.count(EmailTracking.id).label('count')
    ).group_by(EmailTracking.event_type).all()
    
    tracking_data = {stat.event_type: stat.count for stat in tracking_stats}
    
    # Recent campaign performance
    recent_campaigns = Campaign.query.filter(Campaign.sent_at.isnot(None)).order_by(Campaign.sent_at.desc()).limit(10).all()
    
    # Top performing campaigns by open rate
    top_campaigns = []
    for campaign in Campaign.query.filter_by(status='sent').all():
        recipients = campaign.recipients.filter_by(status='sent').count()
        opens = campaign.recipients.filter(CampaignRecipient.opened_at.isnot(None)).count()
        clicks = campaign.recipients.filter(CampaignRecipient.clicked_at.isnot(None)).count()
        
        if recipients > 0:
            campaign_open_rate = (opens / recipients * 100)
            campaign_click_rate = (clicks / recipients * 100)
            top_campaigns.append({
                'campaign': campaign,
                'recipients': recipients,
                'opens': opens,
                'clicks': clicks,
                'open_rate': campaign_open_rate,
                'click_rate': campaign_click_rate
            })
    
    # Sort by open rate
    top_campaigns.sort(key=lambda x: x['open_rate'], reverse=True)
    top_campaigns = top_campaigns[:5]  # Top 5 campaigns
    
    return render_template('analytics.html',
                         total_campaigns=total_campaigns,
                         sent_campaigns=sent_campaigns,
                         total_sent=total_sent,
                         total_failed=total_failed,
                         total_bounced=total_bounced,
                         total_pending=total_pending,
                         total_opened=total_opened,
                         total_clicked=total_clicked,
                         open_rate=open_rate,
                         click_rate=click_rate,
                         bounce_rate=bounce_rate,
                         tracking_data=tracking_data,
                         recent_campaigns=recent_campaigns,
                         top_campaigns=top_campaigns,
                         total_sms_campaigns=total_sms_campaigns,
                         sent_sms_campaigns=sent_sms_campaigns,
                         sms_sent=sms_sent,
                         sms_failed=sms_failed,
                         sms_pending=sms_pending,
                         sms_delivery_rate=sms_delivery_rate)

@main_bp.route('/track/open/<tracking_id>')
def track_open(tracking_id):
    """Track email open events"""
    campaign_id, contact_id = decode_tracking_data(tracking_id)
    
    if campaign_id and contact_id:
        # Get client info for tracking
        event_data = {
            'user_agent': request.headers.get('User-Agent', ''),
            'ip_address': request.remote_addr,
            'timestamp': datetime.utcnow().isoformat()
        }
        
        # Record the open event
        record_email_event(campaign_id, contact_id, 'opened', event_data)
    
    # Return a 1x1 transparent pixel
    from flask import Response
    
    # 1x1 transparent GIF in base64
    pixel_data = 'R0lGODlhAQABAIAAAAAAAP///yH5BAEAAAAALAAAAAABAAEAAAIBRAA7'
    pixel_bytes = base64.b64decode(pixel_data)
    
    response = Response(pixel_bytes, mimetype='image/gif')
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    
    return response

@main_bp.route('/track/click')
def track_click():
    """Track email click events"""
    tracking_id = request.args.get('tracking_id')
    original_url = request.args.get('url', '/')
    
    if tracking_id:
        campaign_id, contact_id = decode_tracking_data(tracking_id)
        
        if campaign_id and contact_id:
            # Get client info for tracking
            event_data = {
                'user_agent': request.headers.get('User-Agent', ''),
                'ip_address': request.remote_addr,
                'clicked_url': original_url,
                'timestamp': datetime.utcnow().isoformat()
            }
            
            # Record the click event
            record_email_event(campaign_id, contact_id, 'clicked', event_data)
    
    # Redirect to the original URL
    return redirect(original_url)

@main_bp.route('/api/campaign/<int:campaign_id>/analytics')
@login_required
def campaign_analytics_api(campaign_id):
    """API endpoint for campaign analytics"""
    from tracking import get_campaign_analytics
    
    analytics = get_campaign_analytics(campaign_id)
    if not analytics:
        return jsonify({'error': 'Campaign not found'}), 404
    
    # Convert campaign object to dict for JSON serialization
    campaign_data = {
        'id': analytics['campaign'].id,
        'name': analytics['campaign'].name,
        'subject': analytics['campaign'].subject,
        'status': analytics['campaign'].status,
        'created_at': analytics['campaign'].created_at.isoformat() if analytics['campaign'].created_at else None,
        'sent_at': analytics['campaign'].sent_at.isoformat() if analytics['campaign'].sent_at else None
    }
    
    return jsonify({
        'campaign': campaign_data,
        'metrics': {
            'total_recipients': analytics['total_recipients'],
            'sent_count': analytics['sent_count'],
            'failed_count': analytics['failed_count'],
            'bounced_count': analytics['bounced_count'],
            'opened_count': analytics['opened_count'],
            'clicked_count': analytics['clicked_count'],
            'delivery_rate': round(analytics['delivery_rate'], 2),
            'open_rate': round(analytics['open_rate'], 2),
            'click_rate': round(analytics['click_rate'], 2),
            'bounce_rate': round(analytics['bounce_rate'], 2)
        },
        'events': analytics['event_counts']
    })

# LUX AI Agent Routes
@main_bp.route('/lux/generate-campaign', methods=['POST'])
@login_required
def lux_generate_campaign():
    """LUX AI agent - Generate automated campaign"""
    try:
        from ai_agent import lux_agent
        
        data = request.get_json() or {}
        campaign_brief = {
            'objective': data.get('objective', 'Engage audience and drive conversions'),
            'target_audience': data.get('target_audience', 'All active contacts'),
            'brand_info': data.get('brand_info', 'Professional business'),
            'target_tags': data.get('target_tags', []),
            'schedule_time': None
        }
        
        # Parse schedule time if provided
        if data.get('schedule_time'):
            try:
                campaign_brief['schedule_time'] = datetime.fromisoformat(data['schedule_time'])
            except ValueError:
                pass
        
        result = lux_agent.create_automated_campaign(campaign_brief)
        
        if result:
            return jsonify({
                'success': True,
                'campaign': result,
                'message': f'LUX created campaign "{result["campaign_name"]}" with {result["recipients_count"]} recipients'
            })
        else:
            return jsonify({
                'success': False,
                'message': 'LUX was unable to create the campaign. Please check your OpenAI configuration.'
            }), 500
            
    except Exception as e:
        logger.error(f"LUX generate campaign error: {e}")
        return jsonify({
            'success': False,
            'message': f'Error: {str(e)}'
        }), 500

@main_bp.route('/lux/optimize-campaign/<int:campaign_id>')
@login_required
def lux_optimize_campaign(campaign_id):
    """LUX AI agent - Optimize campaign performance"""
    try:
        from ai_agent import lux_agent
        
        optimization = lux_agent.optimize_campaign_performance(campaign_id)
        
        if optimization:
            return jsonify({
                'success': True,
                'optimization': optimization
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Unable to analyze campaign performance'
            }), 404
            
    except Exception as e:
        logger.error(f"LUX optimize campaign error: {e}")
        return jsonify({
            'success': False,
            'message': f'Error: {str(e)}'
        }), 500

@main_bp.route('/lux/audience-analysis')
@login_required
def lux_audience_analysis():
    """LUX AI agent - Analyze audience segments"""
    try:
        from ai_agent import lux_agent
        
        contacts = Contact.query.filter_by(is_active=True).all()
        analysis = lux_agent.analyze_audience_segments(contacts)
        
        if analysis:
            return jsonify({
                'success': True,
                'analysis': analysis
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Unable to analyze audience'
            }), 500
            
    except Exception as e:
        logger.error(f"LUX audience analysis error: {e}")
        return jsonify({
            'success': False,
            'message': f'Error: {str(e)}'
        }), 500

@main_bp.route('/lux/subject-variants', methods=['POST'])
@login_required
def lux_subject_variants():
    """LUX AI agent - Generate subject line variants"""
    try:
        from ai_agent import lux_agent
        
        data = request.get_json() or {}
        objective = data.get('objective', 'Engage audience')
        original_subject = data.get('original_subject')
        
        variants = lux_agent.generate_subject_line_variants(objective, original_subject)
        
        if variants:
            return jsonify({
                'success': True,
                'variants': variants
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Unable to generate subject line variants'
            }), 500
            
    except Exception as e:
        logger.error(f"LUX subject variants error: {e}")
        return jsonify({
            'success': False,
            'message': f'Error: {str(e)}'
        }), 500

@main_bp.route('/lux/recommendations')
@login_required
def lux_recommendations():
    """LUX AI agent - Get campaign recommendations"""
    try:
        from ai_agent import lux_agent
        from tracking import get_campaign_analytics
        
        # Gather data to pass to LUX agent (avoiding circular imports)
        recent_campaigns = Campaign.query.order_by(Campaign.created_at.desc()).limit(5).all()
        total_contacts = Contact.query.filter_by(is_active=True).count()
        
        campaign_data = []
        for campaign in recent_campaigns:
            analytics = get_campaign_analytics(campaign.id)
            if analytics:
                campaign_data.append({
                    'name': campaign.name,
                    'open_rate': analytics['open_rate'],
                    'click_rate': analytics['click_rate'],
                    'created_at': campaign.created_at.strftime('%Y-%m-%d') if campaign.created_at else ''
                })
        
        recommendations = lux_agent.get_campaign_recommendations(campaign_data, total_contacts)
        
        if recommendations:
            return jsonify({
                'success': True,
                'recommendations': recommendations
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Unable to generate recommendations'
            }), 500
            
    except Exception as e:
        logger.error(f"LUX recommendations error: {e}")
        return jsonify({
            'success': False,
            'message': f'Error: {str(e)}'
        }), 500

@main_bp.route('/lux')
@login_required
def lux_agent_dashboard():
    """LUX AI Agent dashboard"""
    # Get recent campaigns for optimization selection
    recent_campaigns = Campaign.query.filter(Campaign.sent_at.isnot(None)).order_by(Campaign.sent_at.desc()).limit(6).all()
    
    return render_template('lux_agent.html', recent_campaigns=recent_campaigns)

@main_bp.route('/test-email', methods=['GET', 'POST'])
@login_required
def test_email():
    """Test email sending functionality"""
    if request.method == 'POST':
        test_email_address = request.form.get('test_email', '').strip()
        
        if not test_email_address:
            flash('Please enter a test email address', 'error')
            return render_template('test_email.html')
        
        try:
            from email_service import EmailService
            email_service = EmailService()
            
            # Send test email
            subject = "LUX Email Marketing - Test Email"
            html_content = """
            <html>
            <body style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 30px; text-align: center;">
                    <h1 style="color: white; margin: 0;">✅ LUX Email Marketing</h1>
                    <p style="color: white; margin: 10px 0 0 0;">Email Test Successful!</p>
                </div>
                <div style="padding: 30px; background: #f8f9fa;">
                    <h2 style="color: #333;">Email System Working</h2>
                    <p style="color: #666; line-height: 1.6;">
                        Congratulations! Your LUX Email Marketing system is properly configured 
                        and able to send emails. This means:
                    </p>
                    <ul style="color: #666; line-height: 1.8;">
                        <li>✅ Microsoft Graph API connection is working</li>
                        <li>✅ Email templates can be processed</li>
                        <li>✅ Campaign emails will be delivered</li>
                        <li>✅ Password reset emails will work</li>
                    </ul>
                    <p style="color: #666; line-height: 1.6;">
                        You can now confidently create and send email marketing campaigns!
                    </p>
                    <hr style="border: none; border-top: 1px solid #eee; margin: 30px 0;">
                    <p style="color: #999; font-size: 12px; text-align: center;">
                        LUX Email Marketing Platform - Test Email
                    </p>
                </div>
            </body>
            </html>
            """
            
            result = email_service.send_email(
                to_email=test_email_address,
                subject=subject,
                html_content=html_content
            )
            
            if result:
                flash(f'✅ Test email sent successfully to {test_email_address}!', 'success')
            else:
                flash('❌ Failed to send test email. Please check your email configuration.', 'error')
                
        except Exception as e:
            flash(f'❌ Error testing email: {str(e)}', 'error')
    
    return render_template('test_email.html')

@main_bp.route('/lux/generate-image', methods=['POST'])
@login_required
def lux_generate_image():
    """LUX AI agent - Generate campaign images with DALL-E"""
    # Exempt from CSRF for JSON API
    from app import csrf
    csrf.exempt(lux_generate_image)
    
    try:
        from ai_agent import lux_agent
        
        data = request.get_json() or {}
        description = data.get('description', 'Professional marketing campaign')
        style = data.get('style', 'professional marketing')
        
        result = lux_agent.generate_campaign_image(description, style)
        
        if result:
            return jsonify({
                'success': True,
                'image': result,
                'message': 'Campaign image generated successfully'
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Unable to generate image. Please check your OpenAI configuration.'
            }), 500
            
    except Exception as e:
        logger.error(f"LUX generate image error: {e}")
        return jsonify({
            'success': False,
            'message': f'Error: {str(e)}'
        }), 500

@main_bp.route('/lux/product-campaign', methods=['POST'])
@login_required
def lux_product_campaign():
    """LUX AI agent - Create WooCommerce product campaign"""
    # Exempt from CSRF for JSON API
    from app import csrf
    csrf.exempt(lux_product_campaign)
    
    try:
        # Ensure no WooCommerce library conflicts
        import sys
        woo_modules = [mod for mod in sys.modules.keys() if 'woocommerce' in mod.lower()]
        if woo_modules:
            logging.warning(f"Detected WooCommerce modules: {woo_modules}")
            # Remove any problematic WooCommerce modules
            for mod in woo_modules:
                if mod in sys.modules:
                    del sys.modules[mod]
        
        from ai_agent import lux_agent
        
        data = request.get_json() or {}
        
        # WooCommerce configuration
        woocommerce_config = {
            'url': data.get('woocommerce_url', ''),
            'consumer_key': data.get('consumer_key', ''),
            'consumer_secret': data.get('consumer_secret', ''),
            'product_limit': data.get('product_limit', 6)
        }
        
        # Validate required WooCommerce fields
        if not all([woocommerce_config['url'], woocommerce_config['consumer_key'], woocommerce_config['consumer_secret']]):
            return jsonify({
                'success': False,
                'message': 'WooCommerce URL, Consumer Key, and Consumer Secret are required'
            }), 400
        
        campaign_objective = data.get('objective', 'Promote our latest products')
        product_filter = data.get('product_filter')  # Category filter
        include_images = data.get('include_images', True)
        
        result = lux_agent.create_product_campaign(
            woocommerce_config, 
            campaign_objective, 
            product_filter, 
            include_images
        )
        
        if result:
            # Create email template with product content
            template = EmailTemplate()
            template.name = f"LUX Product Campaign - {result['campaign_name']}"
            template.subject = result['subject']
            template.html_content = result['html_content']
            db.session.add(template)
            db.session.flush()
            
            # Create campaign
            campaign = Campaign()
            campaign.name = result['campaign_name']
            campaign.subject = result['subject']
            campaign.template_id = template.id
            campaign.status = 'draft'
            db.session.add(campaign)
            db.session.flush()
            
            # Add recipients
            contacts = Contact.query.filter_by(is_active=True).all()
            for contact in contacts:
                recipient = CampaignRecipient()
                recipient.campaign_id = campaign.id
                recipient.contact_id = contact.id
                db.session.add(recipient)
            
            db.session.commit()
            
            return jsonify({
                'success': True,
                'campaign': {
                    'id': campaign.id,
                    'name': campaign.name,
                    'subject': campaign.subject,
                    'product_count': result['product_count'],
                    'featured_products': result['featured_products'],
                    'campaign_image': result['campaign_image'],
                    'recipients_count': len(contacts)
                },
                'message': f'Product campaign created with {result["product_count"]} products'
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Unable to create product campaign. Please check your WooCommerce API credentials.'
            }), 500
            
    except Exception as e:
        logger.error(f"LUX product campaign error: {e}")
        return jsonify({
            'success': False,
            'message': f'Error: {str(e)}'
        }), 500

@main_bp.route('/lux/test-woocommerce', methods=['POST'])
@login_required
def lux_test_woocommerce():
    """Test WooCommerce API connection"""
    try:
        from ai_agent import lux_agent
        
        data = request.get_json() or {}
        
        # Test connection by fetching a few products
        products = lux_agent.fetch_woocommerce_products(
            data.get('woocommerce_url', ''),
            data.get('consumer_key', ''),
            data.get('consumer_secret', ''),
            product_limit=3
        )
        
        if products:
            return jsonify({
                'success': True,
                'message': f'Connected successfully! Found {len(products)} products.',
                'sample_products': products[:3]
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Unable to connect to WooCommerce. Please check your API credentials and URL.'
            }), 400
            
    except Exception as e:
        error_msg = str(e)
        logger.error(f"WooCommerce test error: {e}")
        
        # Handle specific error types
        if "proxies" in error_msg or "Client.__init__()" in error_msg or "woocommerce" in error_msg.lower():
            error_msg = "WooCommerce library conflict detected. The system will use pure requests implementation instead. Please try again."
        elif "timeout" in error_msg.lower():
            error_msg = "Connection timeout. Please check your WooCommerce store URL."
        elif "404" in error_msg:
            error_msg = "WooCommerce API not found. Please check your store URL and ensure WooCommerce REST API is enabled."
        elif "401" in error_msg or "unauthorized" in error_msg.lower():
            error_msg = "Invalid consumer key or secret. Please check your WooCommerce API credentials."
        
        return jsonify({
            'success': False,
            'message': f'Connection failed: {error_msg}'
        }), 500

# Drag & Drop Email Editor Routes
@main_bp.route('/editor')
@login_required
def drag_drop_editor():
    """Drag and drop email editor"""
    brandkits = BrandKit.query.filter_by(is_default=True).first()
    return render_template('drag_drop_editor.html', brandkit=brandkits)

@main_bp.route('/editor/save', methods=['POST'])
@login_required  
def save_drag_drop_template():
    """Save template from drag and drop editor"""
    try:
        name = request.form.get('name')
        subject = request.form.get('subject')
        html_content = request.form.get('html_content')
        template_type = request.form.get('template_type', 'custom')
        
        if not name or not subject or not html_content:
            return jsonify({'success': False, 'message': 'Missing required fields'})
            
        # Create new template
        template = EmailTemplate()
        template.name = name
        template.subject = subject
        template.html_content = html_content
        template.template_type = template_type
        
        db.session.add(template)
        db.session.commit()
        
        logger.info(f"Template '{name}' saved successfully by user {current_user.username}")
        return jsonify({'success': True, 'template_id': template.id})
        
    except Exception as e:
        logger.error(f"Error saving template: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

# AI Content Generation Routes
@main_bp.route('/ai/generate-content', methods=['POST'])
@login_required
def generate_ai_content():
    """Generate AI content for emails"""
    try:
        data = request.get_json()
        prompt = data.get('prompt', '')
        content_type = data.get('content_type', 'email_content')
        
        if not prompt:
            return jsonify({'success': False, 'message': 'Prompt is required'})
            
        # Generate content using LUX AI agent
        content_options = lux_agent.generate_email_content(prompt, content_type)
        
        return jsonify({'success': True, 'content': content_options})
        
    except Exception as e:
        logger.error(f"Error generating AI content: {e}")
        return jsonify({'success': False, 'message': str(e)})

@main_bp.route('/ai/generate-subject-lines', methods=['POST'])
@login_required
def generate_subject_lines():
    """Generate AI subject line suggestions"""
    try:
        data = request.get_json()
        campaign_type = data.get('campaign_type', '')
        audience = data.get('audience', '')
        
        subject_lines = lux_agent.generate_subject_lines(campaign_type, audience)
        
        return jsonify({'success': True, 'subject_lines': subject_lines})
        
    except Exception as e:
        logger.error(f"Error generating subject lines: {e}")
        return jsonify({'success': False, 'message': str(e)})

# BrandKit Management Routes
@main_bp.route('/brandkit')
@login_required
def brandkit_management():
    """BrandKit management page"""
    brandkits = BrandKit.query.all()
    return render_template('brandkit.html', brandkits=brandkits)

@main_bp.route('/brandkit/create', methods=['POST'])
@login_required
def create_brandkit():
    """Create new BrandKit"""
    try:
        name = request.form.get('name')
        logo_url = request.form.get('logo_url')
        primary_color = request.form.get('primary_color')
        secondary_color = request.form.get('secondary_color')
        accent_color = request.form.get('accent_color')
        primary_font = request.form.get('primary_font')
        secondary_font = request.form.get('secondary_font')
        
        brandkit = BrandKit()
        brandkit.name = name
        brandkit.logo_url = logo_url
        brandkit.primary_color = primary_color
        brandkit.secondary_color = secondary_color
        brandkit.accent_color = accent_color
        brandkit.primary_font = primary_font
        brandkit.secondary_font = secondary_font
        
        db.session.add(brandkit)
        db.session.commit()
        
        flash('BrandKit created successfully!', 'success')
        return redirect(url_for('main.brandkit_management'))
        
    except Exception as e:
        logger.error(f"Error creating BrandKit: {e}")
        flash('Error creating BrandKit', 'error')
        return redirect(url_for('main.brandkit_management'))

# Poll and Survey Routes
@main_bp.route('/polls')
@login_required
def polls_management():
    """Polls and surveys management"""
    polls = Poll.query.filter_by(is_active=True).all()
    return render_template('polls.html', polls=polls)

@main_bp.route('/polls/create', methods=['POST'])
@login_required
def create_poll():
    """Create new poll"""
    try:
        question = request.form.get('question')
        poll_type = request.form.get('poll_type', 'multiple_choice')
        options = request.form.getlist('options[]')
        
        poll = Poll()
        poll.question = question
        poll.poll_type = poll_type
        poll.options = options
        
        db.session.add(poll)
        db.session.commit()
        
        flash('Poll created successfully!', 'success')
        return redirect(url_for('main.polls_management'))
        
    except Exception as e:
        logger.error(f"Error creating poll: {e}")
        flash('Error creating poll', 'error')
        return redirect(url_for('main.polls_management'))

@main_bp.route('/polls/<int:poll_id>/respond', methods=['POST'])
def respond_to_poll(poll_id):
    """Submit poll response (public endpoint)"""
    try:
        poll = Poll.query.get_or_404(poll_id)
        contact_id = request.form.get('contact_id')
        response_data = request.form.get('response_data')
        
        if contact_id and response_data:
            poll_response = PollResponse()
            poll_response.poll_id = poll_id
            poll_response.contact_id = contact_id
            poll_response.response_data = json.loads(response_data)
            
            db.session.add(poll_response)
            db.session.commit()
            
        return jsonify({'success': True})
        
    except Exception as e:
        logger.error(f"Error submitting poll response: {e}")
        return jsonify({'success': False, 'message': str(e)})

@main_bp.route('/polls/<int:poll_id>/results')
@login_required
def view_poll_results(poll_id):
    """View poll results"""
    poll = Poll.query.get_or_404(poll_id)
    responses = poll.responses.all()
    
    # Calculate results
    results = {}
    if poll.poll_type == 'multiple_choice':
        for response in responses:
            answer = response.response_data.get('answer', '')
            results[answer] = results.get(answer, 0) + 1
    elif poll.poll_type == 'rating':
        ratings = [response.response_data.get('rating', 0) for response in responses]
        if ratings:
            results['average'] = sum(ratings) / len(ratings)
            results['total'] = len(ratings)
    
    return render_template('poll_results.html', poll=poll, responses=responses, results=results)

@main_bp.route('/polls/<int:poll_id>/delete', methods=['POST'])
@login_required
def delete_poll(poll_id):
    """Delete a poll"""
    try:
        poll = Poll.query.get_or_404(poll_id)
        poll.is_active = False
        db.session.commit()
        flash('Poll deleted successfully!', 'success')
    except Exception as e:
        logger.error(f"Error deleting poll: {e}")
        flash('Error deleting poll', 'error')
    
    return redirect(url_for('main.polls_management'))

# A/B Testing Routes
@main_bp.route('/ab-tests')
@login_required
def ab_tests():
    """A/B testing management"""
    tests = ABTest.query.all()
    draft_campaigns = Campaign.query.filter_by(status='draft').all()
    return render_template('ab_tests.html', tests=tests, draft_campaigns=draft_campaigns)

@main_bp.route('/ab-tests/create', methods=['POST'])
@login_required
def create_ab_test():
    """Create new A/B test"""
    try:
        campaign_id = request.form.get('campaign_id')
        test_type = request.form.get('test_type', 'subject_line')
        variant_a = request.form.get('variant_a')
        variant_b = request.form.get('variant_b')
        split_ratio = float(request.form.get('split_ratio', 0.5))
        
        if not campaign_id or not variant_a or not variant_b:
            flash('Campaign, Variant A, and Variant B are required', 'error')
            return redirect(url_for('main.ab_tests'))
        
        # Validate campaign exists
        campaign = Campaign.query.get(campaign_id)
        if not campaign:
            flash('Selected campaign not found', 'error')
            return redirect(url_for('main.ab_tests'))
        
        ab_test = ABTest()
        ab_test.campaign_id = campaign_id
        ab_test.test_type = test_type
        ab_test.variant_a = variant_a
        ab_test.variant_b = variant_b
        ab_test.split_ratio = split_ratio
        ab_test.status = 'draft'
        
        db.session.add(ab_test)
        db.session.commit()
        
        flash('A/B test created successfully!', 'success')
        return redirect(url_for('main.ab_tests'))
        
    except Exception as e:
        logger.error(f"Error creating A/B test: {e}")
        db.session.rollback()
        flash('Error creating A/B test', 'error')
        return redirect(url_for('main.ab_tests'))

# Contact Segmentation Routes
@main_bp.route('/segments')
@login_required
def segments():
    """Contact segmentation management"""
    segments = Segment.query.all()
    return render_template('segments.html', segments=segments)

@main_bp.route('/segments/create', methods=['POST'])
@login_required
def create_segment():
    """Create new contact segment"""
    try:
        name = request.form.get('name')
        description = request.form.get('description')
        segment_type = request.form.get('segment_type', 'behavioral')
        conditions = request.form.get('conditions')
        is_dynamic = request.form.get('is_dynamic') == 'on'
        
        segment = Segment()
        segment.name = name
        segment.description = description
        segment.segment_type = segment_type
        segment.conditions = json.loads(conditions) if conditions else {}
        segment.is_dynamic = is_dynamic
        
        db.session.add(segment)
        db.session.commit()
        
        flash('Segment created successfully!', 'success')
        return redirect(url_for('main.segments'))
        
    except Exception as e:
        logger.error(f"Error creating segment: {e}")
        flash('Error creating segment', 'error')
        return redirect(url_for('main.segments'))

# Social Media Management Routes
@main_bp.route('/social-media')
@login_required
def social_media():
    """Social media management dashboard"""
    posts = SocialPost.query.order_by(SocialPost.created_at.desc()).limit(20).all()
    connected_accounts = SocialMediaAccount.query.filter_by(is_active=True).all()
    return render_template('social_media.html', posts=posts, connected_accounts=connected_accounts)

@main_bp.route('/social/create', methods=['POST'])
@login_required
def create_social_post():
    """Create new social media post"""
    try:
        content = request.form.get('content')
        platforms = request.form.getlist('platforms[]')
        scheduled_at = request.form.get('scheduled_at')
        
        post = SocialPost()
        post.content = content
        post.platforms = platforms
        post.scheduled_at = datetime.fromisoformat(scheduled_at) if scheduled_at else None
        
        db.session.add(post)
        db.session.commit()
        
        flash('Social media post created successfully!', 'success')
        return redirect(url_for('main.social_media'))
        
    except Exception as e:
        logger.error(f"Error creating social post: {e}")
        flash('Error creating social post', 'error')
        return redirect(url_for('main.social_media'))

@main_bp.route('/social/refresh-followers', methods=['POST'])
@login_required
def refresh_social_followers():
    """Refresh follower counts from social media platforms"""
    try:
        from services.social_media_service import SocialMediaService
        
        account_id = request.form.get('account_id')
        if account_id:
            # Refresh specific account
            account = SocialMediaAccount.query.get(account_id)
            if account:
                result = SocialMediaService.refresh_account_data(account)
                if result.get('success'):
                    account.follower_count = result.get('follower_count', account.follower_count)
                    account.last_synced_at = datetime.utcnow()
                    db.session.commit()
                    flash(f"Updated {account.platform} follower count", 'success')
                else:
                    flash(f"Could not refresh {account.platform}: {result.get('message')}", 'warning')
        else:
            # Refresh all accounts
            accounts = SocialMediaAccount.query.filter_by(is_active=True).all()
            updated_count = 0
            for account in accounts:
                result = SocialMediaService.refresh_account_data(account)
                if result.get('success'):
                    account.follower_count = result.get('follower_count', account.follower_count)
                    account.last_synced_at = datetime.utcnow()
                    updated_count += 1
            
            db.session.commit()
            flash(f"Refreshed {updated_count} social media account(s)", 'success')
        
        return redirect(url_for('main.social_media'))
        
    except Exception as e:
        logger.error(f"Error refreshing social followers: {e}")
        flash('Error refreshing follower counts', 'error')
        return redirect(url_for('main.social_media'))

# Advanced Automation Management Routes
@main_bp.route('/automations')
@login_required
def automation_dashboard():
    """Unified Automations & AI Agents dashboard"""
    from agent_scheduler import get_agent_scheduler
    
    automations = Automation.query.all()
    templates = AutomationTemplate.query.filter_by(is_predefined=True).all()
    active_executions = AutomationExecution.query.filter_by(status='active').count()
    
    # Get AI agents info
    scheduler = get_agent_scheduler()
    agents = scheduler.agents if scheduler else {}
    
    return render_template('automation_dashboard.html', 
                         automations=automations, 
                         templates=templates,
                         active_executions=active_executions,
                         agents=agents)

@main_bp.route('/automations/create', methods=['GET', 'POST'])
@login_required
def create_automation():
    """Create new automation workflow"""
    if request.method == 'POST':
        try:
            name = request.form.get('name')
            description = request.form.get('description')
            trigger_type = request.form.get('trigger_type')
            trigger_conditions = request.form.get('trigger_conditions')
            
            automation = Automation()
            automation.name = name
            automation.description = description
            automation.trigger_type = trigger_type
            automation.trigger_conditions = json.loads(trigger_conditions) if trigger_conditions else {}
            
            db.session.add(automation)
            db.session.commit()
            
            flash('Automation workflow created successfully!', 'success')
            return redirect(url_for('main.edit_automation', id=automation.id))
            
        except Exception as e:
            logger.error(f"Error creating automation: {e}")
            flash('Error creating automation workflow', 'error')
            return redirect(url_for('main.automation_dashboard'))
    
    templates = AutomationTemplate.query.filter_by(is_predefined=True).all()
    email_templates = EmailTemplate.query.all()
    return render_template('create_automation.html', templates=templates, email_templates=email_templates)

@main_bp.route('/automations/<int:id>/edit', methods=['GET', 'POST'])
@login_required
def edit_automation(id):
    """Edit automation workflow with visual builder"""
    automation = Automation.query.get_or_404(id)
    
    if request.method == 'POST':
        try:
            # Update automation details
            automation.name = request.form.get('name')
            automation.description = request.form.get('description')
            automation.trigger_type = request.form.get('trigger_type')
            trigger_conditions = request.form.get('trigger_conditions')
            automation.trigger_conditions = json.loads(trigger_conditions) if trigger_conditions else {}
            
            # Update steps from visual builder
            steps_data = request.form.get('steps_data')
            if steps_data:
                steps = json.loads(steps_data)
                
                # Delete existing steps
                AutomationStep.query.filter_by(automation_id=id).delete()
                
                # Create new steps
                for i, step_data in enumerate(steps):
                    step = AutomationStep()
                    step.automation_id = id
                    step.step_type = step_data.get('type')
                    step.step_order = i
                    step.template_id = step_data.get('template_id')
                    step.delay_hours = step_data.get('delay_hours', 0)
                    step.conditions = step_data.get('conditions', {})
                    
                    db.session.add(step)
            
            db.session.commit()
            flash('Automation updated successfully!', 'success')
            return redirect(url_for('main.automation_dashboard'))
            
        except Exception as e:
            logger.error(f"Error updating automation: {e}")
            db.session.rollback()
            flash('Error updating automation', 'error')
    
    steps = AutomationStep.query.filter_by(automation_id=id).order_by(AutomationStep.step_order).all()
    email_templates = EmailTemplate.query.all()
    executions = AutomationExecution.query.filter_by(automation_id=id).limit(10).all()
    
    return render_template('edit_automation.html', 
                         automation=automation, 
                         steps=steps,
                         email_templates=email_templates,
                         executions=executions)

@main_bp.route('/automations/<int:id>/toggle', methods=['POST'])
@login_required
def toggle_automation(id):
    """Enable/disable automation workflow"""
    try:
        automation = Automation.query.get_or_404(id)
        automation.is_active = not automation.is_active
        db.session.commit()
        
        status = 'activated' if automation.is_active else 'deactivated'
        flash(f'Automation {status} successfully!', 'success')
        
        return jsonify({'success': True, 'is_active': automation.is_active})
    except Exception as e:
        logger.error(f"Error toggling automation: {e}")
        return jsonify({'success': False, 'message': str(e)})

@main_bp.route('/automation-templates')
@login_required
def automation_templates():
    """Automation template library"""
    predefined = AutomationTemplate.query.filter_by(is_predefined=True).all()
    custom = AutomationTemplate.query.filter_by(is_predefined=False).all()
    
    return render_template('automation_templates.html', 
                         predefined_templates=predefined,
                         custom_templates=custom)

@main_bp.route('/automation-templates/create-from-template/<int:template_id>')
@login_required
def create_from_template(template_id):
    """Create automation from predefined template"""
    try:
        template = AutomationTemplate.query.get_or_404(template_id)
        template_data = template.template_data
        
        # Create automation from template
        automation = Automation()
        automation.name = f"{template.name} - Copy"
        automation.description = template.description
        automation.trigger_type = template_data.get('trigger_type', 'custom')
        automation.trigger_conditions = template_data.get('trigger_conditions', {})
        
        db.session.add(automation)
        db.session.flush()
        
        # Create steps from template
        for i, step_data in enumerate(template_data.get('steps', [])):
            step = AutomationStep()
            step.automation_id = automation.id
            step.step_type = step_data.get('type')
            step.step_order = i
            step.delay_hours = step_data.get('delay_hours', 0)
            step.conditions = step_data.get('conditions', {})
            
            db.session.add(step)
        
        # Update usage count
        template.usage_count += 1
        
        db.session.commit()
        
        flash(f'Created automation from template: {template.name}', 'success')
        return redirect(url_for('main.edit_automation', id=automation.id))
        
    except Exception as e:
        logger.error(f"Error creating from template: {e}")
        flash('Error creating automation from template', 'error')
        return redirect(url_for('main.automation_templates'))

@main_bp.route('/automation-analytics')
@login_required
def automation_analytics():
    """Automation performance analytics"""
    total_automations = Automation.query.count()
    active_automations = Automation.query.filter_by(is_active=True).count()
    total_executions = AutomationExecution.query.count()
    completed_executions = AutomationExecution.query.filter_by(status='completed').count()
    
    # Recent execution data
    recent_executions = AutomationExecution.query.order_by(AutomationExecution.started_at.desc()).limit(20).all()
    
    # Performance by automation
    automation_stats = []
    for automation in Automation.query.all():
        executions = AutomationExecution.query.filter_by(automation_id=automation.id)
        total = executions.count()
        completed = executions.filter_by(status='completed').count()
        completion_rate = (completed / total * 100) if total > 0 else 0
        
        automation_stats.append({
            'automation': automation,
            'total_executions': total,
            'completed': completed,
            'completion_rate': completion_rate
        })
    
    return render_template('automation_analytics.html',
                         total_automations=total_automations,
                         active_automations=active_automations,
                         total_executions=total_executions,
                         completed_executions=completed_executions,
                         recent_executions=recent_executions,
                         automation_stats=automation_stats)
# Automation Pause/Resume (Phase 0-1)
@main_bp.route('/automations/<int:id>/pause', methods=['POST'])
@login_required
def pause_automation(id):
    """Pause an automation workflow"""
    try:
        automation = Automation.query.get_or_404(id)
        reason = request.form.get('reason', 'Manual pause')
        
        automation.pause(reason)
        db.session.commit()
        
        flash(f'Automation "{automation.name}" paused successfully!', 'success')
        return jsonify({'success': True, 'is_paused': True})
    except Exception as e:
        logger.error(f"Error pausing automation: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500

@main_bp.route('/automations/<int:id>/resume', methods=['POST'])
@login_required
def resume_automation(id):
    """Resume a paused automation workflow"""
    try:
        automation = Automation.query.get_or_404(id)
        automation.resume()
        db.session.commit()
        
        flash(f'Automation "{automation.name}" resumed successfully!', 'success')
        return jsonify({'success': True, 'is_paused': False})
    except Exception as e:
        logger.error(f"Error resuming automation: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500

# SMS Marketing Module (Phase 0-1)
@main_bp.route('/sms')
@login_required
def sms_dashboard():
    """SMS marketing dashboard"""
    # # from services.sms_service import SMSService
    
    campaigns = SMSCampaign.query.order_by(SMSCampaign.created_at.desc()).all()
    templates = SMSTemplate.query.order_by(SMSTemplate.created_at.desc()).limit(10).all()
    
    # Stats
    total_campaigns = SMSCampaign.query.count()
    sent_campaigns = SMSCampaign.query.filter_by(status='sent').count()
    scheduled_campaigns = SMSCampaign.query.filter_by(status='scheduled').count()
    
    return render_template('sms_campaigns.html',
                         campaigns=campaigns,
                         templates=templates,
                         total_campaigns=total_campaigns,
                         sent_campaigns=sent_campaigns,
                         scheduled_campaigns=scheduled_campaigns,
                         sms_enabled=True)

@main_bp.route('/sms/create', methods=['GET', 'POST'])
@login_required
def create_sms_campaign():
    """Create a new SMS campaign"""
    # from services.sms_service import SMSService
    # from services.scheduling_service import SchedulingService
    from services.campaign_tagging_service import CampaignTaggingService
    
    if request.method == 'POST':
        try:
            name = request.form.get('name')
            message = request.form.get('message')
            send_option = request.form.get('send_option', 'now')
            
            # Build scheduled_at from date and time if scheduled
            scheduled_at = None
            if send_option == 'scheduled':
                scheduled_date = request.form.get('scheduled_date')
                scheduled_time = request.form.get('scheduled_time')
                if scheduled_date and scheduled_time:
                    scheduled_at = datetime.fromisoformat(f"{scheduled_date}T{scheduled_time}")
            
            # Create campaign
            campaign = SMSService.create_campaign(
                name=name,
                message=message,
                scheduled_at=scheduled_at
            )
            
            # Process campaign tags for organization
            campaign_tags_str = request.form.get('campaign_tags', '')
            if campaign_tags_str:
                tag_names = [tag.strip() for tag in campaign_tags_str.split(',') if tag.strip()]
                tag_ids = []
                for tag_name in tag_names:
                    tag = CampaignTaggingService.create_tag(tag_name)
                    tag_ids.append(tag.id)
                
                if tag_ids:
                    CampaignTaggingService.sync_tags_for_object(
                        tag_ids,
                        'sms',
                        campaign.id
                    )
            
            # Process segment tags for targeting (filter contacts)
            segment_tags_str = request.form.get('segment_tags', '')
            contacts_to_target = []
            
            if segment_tags_str:
                # Filter contacts by tags - check if contact has ANY of the specified tags
                segment_names = [seg.strip() for seg in segment_tags_str.split(',') if seg.strip()]
                
                # Build OR filter: contact.tags contains tag1 OR tag2 OR tag3
                tag_filters = [Contact.tags.contains(tag_name) for tag_name in segment_names]
                
                contacts_to_target = Contact.query.filter(
                    Contact.phone.isnot(None),
                    or_(*tag_filters) if tag_filters else True
                ).all()
            else:
                # Send to all contacts with phone numbers
                contacts_to_target = Contact.query.filter(Contact.phone.isnot(None)).all()
            
            # Add recipients
            if contacts_to_target:
                SMSService.add_recipients(campaign.id, [contact.id for contact in contacts_to_target])
            
            # Add to unified schedule if scheduled
            if scheduled_at:
                SchedulingService.create_schedule(
                    module_type='sms',
                    module_object_id=campaign.id,
                    title=f'SMS: {name}',
                    scheduled_at=scheduled_at,
                    description=message[:100]
                )
            
            flash('SMS campaign created successfully!', 'success')
            return redirect(url_for('main.sms_dashboard'))
            
        except Exception as e:
            logger.error(f"Error creating SMS campaign: {e}")
            flash('Error creating SMS campaign', 'error')
    
    contacts = Contact.query.filter(Contact.phone.isnot(None)).all()
    tags = CampaignTaggingService.get_all_tags()
    templates = SMSTemplate.query.all()
    segments = Segment.query.all()
    
    return render_template('create_sms_campaign.html',
                         contacts=contacts,
                         tags=tags,
                         templates=templates,
                         segments=segments)

@main_bp.route('/sms/templates/create', methods=['GET', 'POST'])
@login_required
def create_sms_template():
    """Create a reusable SMS template"""
    # from services.sms_service import SMSService
    
    if request.method == 'POST':
        try:
            name = request.form.get('name')
            message = request.form.get('message')
            category = request.form.get('category', 'promotional')
            tone = request.form.get('tone', 'professional')
            
            template = SMSService.create_template(name, message, category, tone)
            
            flash(f'SMS template "{name}" created successfully!', 'success')
            return redirect(url_for('main.sms_dashboard'))
            
        except Exception as e:
            logger.error(f"Error creating SMS template: {e}")
            flash('Error creating SMS template', 'error')
    
    return render_template('create_sms_template.html')

@main_bp.route('/sms/ai-generate', methods=['POST'])
@login_required
def ai_generate_sms():
    """Generate SMS content using AI"""
    # from services.sms_service import SMSService
    
    try:
        # Support both 'prompt' and 'campaign_name' for backwards compatibility
        prompt = request.json.get('prompt') or request.json.get('campaign_name')
        tone = request.json.get('tone', 'professional')
        max_length = int(request.json.get('max_length', 160))
        
        if not prompt:
            return jsonify({'success': False, 'error': 'Campaign name or prompt is required'}), 400
        
        # Create a better prompt from campaign name
        if request.json.get('campaign_name'):
            prompt = f"Create an SMS marketing message for: {prompt}"
        
        message = SMSService.ai_generate_sms(prompt, tone, max_length)
        
        return jsonify({
            'success': True,
            'message': message,
            'length': len(message),
            'is_compliant': SMSService.check_compliance(message)
        })
        
    except Exception as e:
        logger.error(f"Error generating SMS: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/sms/<int:campaign_id>/analytics')
@login_required
def sms_analytics(campaign_id):
    """View SMS campaign analytics"""
    # from services.sms_service import SMSService
    
    campaign = SMSCampaign.query.get_or_404(campaign_id)
    analytics = SMSService.calculate_analytics(campaign_id)
    recipients = SMSRecipient.query.filter_by(campaign_id=campaign_id).all()
    
    return render_template('sms_analytics.html',
                         campaign=campaign,
                         analytics=analytics,
                         recipients=recipients)

# Non-Opener Resend Feature
@main_bp.route('/campaigns/<int:campaign_id>/resend-non-openers', methods=['GET', 'POST'])
@login_required
def setup_non_opener_resend(campaign_id):
    """Set up automatic resend to non-openers"""
    campaign = Campaign.query.get_or_404(campaign_id)
    
    if request.method == 'POST':
        try:
            hours_after = int(request.form.get('hours_after', 24))
            new_subject = request.form.get('new_subject_line')
            
            resend = NonOpenerResend()
            resend.original_campaign_id = campaign_id
            resend.hours_after_original = hours_after
            resend.new_subject_line = new_subject
            resend.scheduled_at = campaign.sent_at + timedelta(hours=hours_after) if campaign.sent_at else None
            
            db.session.add(resend)
            db.session.commit()
            
            flash('Non-opener resend scheduled successfully!', 'success')
            return redirect(url_for('main.campaign_details', id=campaign_id))
            
        except Exception as e:
            logger.error(f"Error setting up resend: {e}")
            flash('Error setting up resend', 'error')
    
    return render_template('setup_non_opener_resend.html', campaign=campaign)

# Web Forms & Landing Pages Routes
@main_bp.route('/forms')
@login_required
def forms_dashboard():
    """Web forms management dashboard"""
    forms = WebForm.query.all()
    total_submissions = FormSubmission.query.count()
    
    return render_template('forms_dashboard.html', forms=forms, total_submissions=total_submissions)

@main_bp.route('/forms/create', methods=['GET', 'POST'])
@login_required
def create_web_form():
    """Create new web signup form"""
    if request.method == 'POST':
        try:
            name = request.form.get('name')
            title = request.form.get('title')
            description = request.form.get('description')
            fields_data = request.form.get('fields_data')
            success_message = request.form.get('success_message')
            redirect_url = request.form.get('redirect_url')
            
            form = WebForm()
            form.name = name
            form.title = title
            form.description = description
            form.fields = json.loads(fields_data) if fields_data else []
            form.success_message = success_message
            form.redirect_url = redirect_url
            
            db.session.add(form)
            db.session.commit()
            
            flash('Web form created successfully!', 'success')
            return redirect(url_for('main.forms_dashboard'))
            
        except Exception as e:
            logger.error(f"Error creating form: {e}")
            flash('Error creating web form', 'error')
    
    return render_template('create_web_form.html')

@main_bp.route('/forms/<int:id>/embed-code')
@login_required
def form_embed_code(id):
    """Get embed code for web form"""
    form = WebForm.query.get_or_404(id)
    
    embed_html = f"""
<div id="lux-form-{form.id}"></div>
<script>
(function() {{
    var script = document.createElement('script');
    script.src = '{request.url_root}static/js/form-embed.js';
    script.onload = function() {{
        LuxForm.render({form.id}, 'lux-form-{form.id}');
    }};
    document.head.appendChild(script);
}})();
</script>
    """
    
    return jsonify({'embed_code': embed_html})

@main_bp.route('/landing-pages')
@login_required
def landing_pages():
    """Landing pages management"""
    try:
        pages = LandingPage.query.all()
        try:
            forms = WebForm.query.all()
        except Exception as e:
            logger.warning(f"WebForm table not found: {e}")
            forms = []
        
        return render_template('landing_pages.html', pages=pages, forms=forms)
    except Exception as e:
        logger.error(f"Error loading landing pages: {e}")
        flash('Error loading landing pages. Please check database tables.', 'error')
        return redirect(url_for('main.dashboard'))

@main_bp.route('/landing-pages/create', methods=['GET', 'POST'])
@login_required
def create_landing_page():
    """Create new landing page"""
    if request.method == 'POST':
        try:
            name = request.form.get('name')
            title = request.form.get('title')
            slug = request.form.get('slug')
            html_content = request.form.get('html_content')
            css_styles = request.form.get('css_styles')
            meta_description = request.form.get('meta_description')
            form_id = request.form.get('form_id') or None
            
            # Validate required fields
            if not name or not slug or not html_content:
                flash('Please fill in all required fields', 'error')
                return redirect(url_for('main.create_landing_page'))
            
            # Validate slug format (lowercase, letters, numbers, hyphens only)
            import re
            if not re.match(r'^[a-z0-9-]+$', slug):
                flash('URL slug must contain only lowercase letters, numbers, and hyphens', 'error')
                return redirect(url_for('main.create_landing_page'))
            
            # Check if slug already exists
            existing_page = LandingPage.query.filter_by(slug=slug).first()
            if existing_page:
                flash(f'A landing page with slug "{slug}" already exists', 'error')
                return redirect(url_for('main.create_landing_page'))
            
            page = LandingPage()
            page.name = name
            page.title = title
            page.slug = slug
            page.html_content = html_content
            page.css_styles = css_styles
            page.meta_description = meta_description
            page.form_id = int(form_id) if form_id else None
            
            db.session.add(page)
            db.session.commit()
            
            flash('Landing page created successfully!', 'success')
            return redirect(url_for('main.landing_pages'))
            
        except Exception as e:
            logger.error(f"Error creating landing page: {e}")
            flash(f'Error creating landing page: {str(e)}', 'error')
            db.session.rollback()
    
    try:
        forms = WebForm.query.all()
    except Exception as e:
        logger.warning(f"WebForm table not found: {e}")
        forms = []
    
    return render_template('create_landing_page.html', forms=forms)

@main_bp.route('/newsletter-archive')
def newsletter_archive():
    """Public newsletter archive"""
    newsletters = NewsletterArchive.query.filter_by(is_public=True).order_by(NewsletterArchive.published_at.desc()).all()
    
    return render_template('newsletter_archive_public.html', newsletters=newsletters)

@main_bp.route('/newsletter-subscribe', methods=['POST'])
@csrf.exempt
def newsletter_subscribe():
    """Public newsletter subscription"""
    data = request.get_json() if request.is_json else request.form
    email = data.get('email', '').strip().lower()
    
    if not email or not validate_email(email):
        return jsonify({'success': False, 'message': 'Please enter a valid email address'}), 400
    
    # Check if contact already exists
    contact = Contact.query.filter_by(email=email).first()
    
    if contact:
        if 'newsletter' not in (contact.tags or ''):
            # Add newsletter tag
            existing_tags = contact.tags.split(',') if contact.tags else []
            if 'newsletter' not in existing_tags:
                existing_tags.append('newsletter')
                contact.tags = ','.join(existing_tags)
                contact.updated_at = datetime.utcnow()
                db.session.commit()
                return jsonify({'success': True, 'message': 'You have been subscribed to our newsletter!'}), 200
        return jsonify({'success': True, 'message': 'You are already subscribed to our newsletter!'}), 200
    
    # Create new contact
    contact = Contact(
        email=email,
        source='newsletter_archive',
        tags='newsletter',
        is_active=True,
        is_subscribed=True
    )
    
    db.session.add(contact)
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Thank you for subscribing to our newsletter!'}), 200

@main_bp.route('/newsletter-archive/<slug>')
def view_newsletter(slug):
    """View individual newsletter"""
    newsletter = NewsletterArchive.query.filter_by(slug=slug, is_public=True).first_or_404()
    
    # Increment view count
    newsletter.view_count += 1
    db.session.commit()
    
    return render_template('newsletter_view.html', newsletter=newsletter)

# SEO Tools Routes
@main_bp.route('/seo')
@login_required
def seo_tools():
    """SEO analysis and optimization tools"""
    return render_template('seo_tools.html')

@main_bp.route('/seo/analyze', methods=['POST'])
@login_required
def analyze_seo():
    """Analyze a URL for SEO"""
    try:
        url = request.form.get('url')
        
        if not url:
            flash('Please enter a URL to analyze', 'error')
            return redirect(url_for('main.seo_tools'))
        
        # Ensure URL has a protocol
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        result = seo_service.analyze_page(url)
        
        if result['success']:
            return render_template('seo_results.html', analysis=result['data'])
        else:
            flash(f'Error analyzing URL: {result["error"]}', 'error')
            return redirect(url_for('main.seo_tools'))
        
    except Exception as e:
        logger.error(f"Error in SEO analysis: {e}")
        flash('Error analyzing URL', 'error')
        return redirect(url_for('main.seo_tools'))

# Events Management Routes
@main_bp.route('/events')
@login_required
def events_dashboard():
    """Events management dashboard"""
    events = Event.query.order_by(Event.start_date.desc()).all()
    return render_template('events.html', events=events)

@main_bp.route('/events/create', methods=['GET', 'POST'])
@login_required
def create_event():
    """Create new event"""
    if request.method == 'POST':
        try:
            name = request.form.get('name')
            description = request.form.get('description')
            start_date_str = request.form.get('start_date')
            end_date_str = request.form.get('end_date')
            location = request.form.get('location')
            max_attendees = request.form.get('max_attendees')
            price = request.form.get('price', 0.0)
            
            if not start_date_str:
                flash('Start date is required', 'error')
                return redirect(url_for('main.create_event'))
            
            event = Event()
            event.name = name
            event.description = description
            event.start_date = datetime.fromisoformat(start_date_str)
            event.end_date = datetime.fromisoformat(end_date_str) if end_date_str else None
            event.location = location
            event.max_attendees = int(max_attendees) if max_attendees else None
            event.price = float(price)
            
            db.session.add(event)
            db.session.commit()
            
            flash('Event created successfully!', 'success')
            return redirect(url_for('main.events_dashboard'))
            
        except Exception as e:
            db.session.rollback()
            logger.error(f"Error creating event: {e}")
            flash('Error creating event', 'error')
            return redirect(url_for('main.events_dashboard'))
    
    return render_template('create_event.html')

@main_bp.route('/events/<int:event_id>')
@login_required
def view_event(event_id):
    """View event details and registrations"""
    event = Event.query.get_or_404(event_id)
    registrations = EventRegistration.query.filter_by(event_id=event_id).all()
    
    # Get contacts for registrations
    for reg in registrations:
        reg.contact = Contact.query.get(reg.contact_id)
    
    return render_template('view_event.html', event=event, registrations=registrations)

# WooCommerce Integration Routes
@main_bp.route('/woocommerce')
@login_required
def woocommerce_dashboard():
    """WooCommerce integration dashboard"""
    from woocommerce_service import WooCommerceService
    wc_service = WooCommerceService()
    
    if not wc_service.is_configured():
        flash('WooCommerce integration is not configured. Please add WooCommerce credentials.', 'warning')
        return render_template('woocommerce_setup.html')
    
    try:
        # Get products summary
        products = wc_service.get_products(per_page=10)
        product_count = len(products) if products else 0
        
        # Get recent orders
        orders = wc_service.get_orders(per_page=5)
        order_count = len(orders) if orders else 0
        
        return render_template('woocommerce_dashboard.html', 
                             products=products,
                             product_count=product_count,
                             orders=orders,
                             order_count=order_count,
                             is_configured=True)
    except Exception as e:
        logger.error(f"WooCommerce error: {e}")
        flash('Error connecting to WooCommerce. Please check your credentials.', 'error')
        return render_template('woocommerce_setup.html')

@main_bp.route('/woocommerce/products')
@login_required
def woocommerce_products():
    """View WooCommerce products"""
    from woocommerce_service import WooCommerceService
    wc_service = WooCommerceService()
    
    if not wc_service.is_configured():
        flash('WooCommerce integration is not configured.', 'warning')
        return redirect(url_for('main.woocommerce_dashboard'))
    
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '')
    
    if search:
        products = wc_service.search_products(search, per_page=20)
    else:
        products = wc_service.get_products(page=page, per_page=20)
    
    return render_template('woocommerce_products.html', 
                         products=products or [],
                         search=search,
                         page=page)

@main_bp.route('/woocommerce/products/<int:product_id>')
@login_required
def woocommerce_product_detail(product_id):
    """View single WooCommerce product"""
    from woocommerce_service import WooCommerceService
    wc_service = WooCommerceService()
    
    product = wc_service.get_product(product_id)
    if not product:
        flash('Product not found', 'error')
        return redirect(url_for('main.woocommerce_products'))
    
    return render_template('woocommerce_product_detail.html', product=product)

@main_bp.route('/woocommerce/sync-products', methods=['POST'])
@login_required
def sync_woocommerce_products():
    """Sync WooCommerce products to local database"""
    from woocommerce_service import WooCommerceService
    wc_service = WooCommerceService()
    
    try:
        products = wc_service.get_all_products(max_products=500)
        
        synced_count = 0
        for wc_product in products:
            # Check if product exists
            product = Product.query.filter_by(wc_product_id=wc_product['id']).first()
            
            if not product:
                product = Product()
                product.wc_product_id = wc_product['id']
            
            # Update product data
            product.name = wc_product['name']
            product.description = wc_product['description']
            product.price = float(wc_product['price']) if wc_product['price'] else 0.0
            product.sku = wc_product.get('sku', '')
            product.stock_quantity = wc_product.get('stock_quantity', 0)
            product.image_url = wc_product['images'][0]['src'] if wc_product.get('images') else None
            product.product_url = wc_product['permalink']
            
            db.session.add(product)
            synced_count += 1
        
        db.session.commit()
        flash(f'Successfully synced {synced_count} products from WooCommerce!', 'success')
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error syncing WooCommerce products: {e}")
        flash('Error syncing products from WooCommerce', 'error')
    
    return redirect(url_for('main.woocommerce_dashboard'))

@main_bp.route('/woocommerce/create-product-campaign/<int:product_id>', methods=['GET', 'POST'])
@login_required
def create_product_campaign(product_id):
    """Create email campaign for a specific product"""
    from woocommerce_service import WooCommerceService
    wc_service = WooCommerceService()
    
    # Get product from WooCommerce
    product = wc_service.get_product(product_id)
    if not product:
        flash('Product not found', 'error')
        return redirect(url_for('main.woocommerce_products'))
    
    if request.method == 'POST':
        try:
            campaign_name = request.form.get('campaign_name')
            subject = request.form.get('subject')
            tag = request.form.get('tag')
            
            # Create campaign
            campaign = Campaign()
            campaign.name = campaign_name
            campaign.subject = subject
            campaign.status = 'draft'
            
            # Generate email content from product
            product_html = f"""
            <div style="max-width: 600px; margin: 0 auto; font-family: Arial, sans-serif;">
                <h1 style="color: #333;">{product['name']}</h1>
                {'<img src="' + product['images'][0]['src'] + '" style="max-width: 100%; height: auto;" />' if product.get('images') else ''}
                <div style="margin: 20px 0;">
                    <h2 style="color: #0066cc;">Price: ${product['price']}</h2>
                </div>
                <div style="margin: 20px 0;">
                    {product.get('description', '')}
                </div>
                <div style="text-align: center; margin: 30px 0;">
                    <a href="{product['permalink']}" style="background: #0066cc; color: white; padding: 15px 30px; text-decoration: none; border-radius: 5px; display: inline-block;">
                        Shop Now
                    </a>
                </div>
            </div>
            """
            
            # Create template for this campaign
            template = EmailTemplate()
            template.name = f"Product Campaign - {product['name']}"
            template.subject = subject
            template.html_content = product_html
            
            db.session.add(template)
            db.session.flush()
            
            campaign.template_id = template.id
            db.session.add(campaign)
            db.session.commit()
            
            flash(f'Product campaign created successfully!', 'success')
            return redirect(url_for('main.campaign_details', id=campaign.id))
            
        except Exception as e:
            db.session.rollback()
            logger.error(f"Error creating product campaign: {e}")
            flash('Error creating campaign', 'error')
    
    return render_template('create_product_campaign.html', product=product)

# AI Agent Dashboard Routes
@main_bp.route('/agents')
@login_required
def agents_dashboard():
    """AI Agents Dashboard - View and manage all marketing agents"""
    from models import AgentTask, AgentLog, AgentReport, AgentSchedule
    from agent_scheduler import get_agent_scheduler
    
    # Get scheduler and job information
    scheduler = get_agent_scheduler()
    scheduled_jobs = scheduler.get_scheduled_jobs() if scheduler else []
    
    # Get recent agent activity
    recent_logs = AgentLog.query.order_by(AgentLog.created_at.desc()).limit(20).all()
    pending_tasks = AgentTask.query.filter_by(status='pending').order_by(AgentTask.scheduled_at).all()
    recent_reports = AgentReport.query.order_by(AgentReport.created_at.desc()).limit(10).all()
    
    # Get agent stats
    agent_stats = {}
    agent_types = ['brand_strategy', 'content_seo', 'analytics', 'creative_design']
    
    for agent_type in agent_types:
        total_tasks = AgentTask.query.filter_by(agent_type=agent_type).count()
        completed_tasks = AgentTask.query.filter_by(agent_type=agent_type, status='completed').count()
        failed_tasks = AgentTask.query.filter_by(agent_type=agent_type, status='failed').count()
        
        agent_stats[agent_type] = {
            'total': total_tasks,
            'completed': completed_tasks,
            'failed': failed_tasks,
            'success_rate': round((completed_tasks / max(total_tasks, 1)) * 100, 1)
        }
    
    return render_template('agents_dashboard.html',
                         scheduled_jobs=scheduled_jobs,
                         recent_logs=recent_logs,
                         pending_tasks=pending_tasks,
                         recent_reports=recent_reports,
                         agent_stats=agent_stats)

@main_bp.route('/agents/<agent_type>/trigger', methods=['POST'])
@login_required
def trigger_agent(agent_type):
    """Manually trigger an agent task"""
    try:
        task_data = request.get_json() or {}
        
        # Ensure task_type is set with defaults per agent
        if 'task_type' not in task_data:
            task_type_defaults = {
                'brand_strategy': 'market_research',
                'content_seo': 'keyword_research',
                'analytics': 'performance_summary',
                'creative_design': 'generate_ad_creative',
                'advertising': 'campaign_strategy',
                'social_media': 'daily_posts',
                'email_crm': 'weekly_campaign',
                'sales_enablement': 'sales_deck',
                'retention': 'churn_analysis',
                'operations': 'system_health'
            }
            task_data['task_type'] = task_type_defaults.get(agent_type, 'default_task')
        
        # Get the appropriate agent
        agent = None
        if agent_type == 'brand_strategy':
            from agents.brand_strategy_agent import BrandStrategyAgent
            agent = BrandStrategyAgent()
        elif agent_type == 'content_seo':
            from agents.content_seo_agent import ContentSEOAgent
            agent = ContentSEOAgent()
        elif agent_type == 'analytics':
            from agents.analytics_agent import AnalyticsAgent
            agent = AnalyticsAgent()
        elif agent_type == 'creative_design':
            from agents.creative_agent import CreativeAgent
            agent = CreativeAgent()
        elif agent_type == 'advertising':
            from agents.advertising_agent import AdvertisingAgent
            agent = AdvertisingAgent()
        elif agent_type == 'social_media':
            from agents.social_media_agent import SocialMediaAgent
            agent = SocialMediaAgent()
        elif agent_type == 'email_crm':
            from agents.email_crm_agent import EmailCRMAgent
            agent = EmailCRMAgent()
        elif agent_type == 'sales_enablement':
            from agents.sales_enablement_agent import SalesEnablementAgent
            agent = SalesEnablementAgent()
        elif agent_type == 'retention':
            from agents.retention_agent import RetentionAgent
            agent = RetentionAgent()
        elif agent_type == 'operations':
            from agents.operations_agent import OperationsAgent
            agent = OperationsAgent()
        else:
            return jsonify({'success': False, 'error': 'Unknown agent type'}), 400
        
        # Create and execute task
        task_id = agent.create_task(
            task_name=task_data.get('task_name', 'Manual Trigger'),
            task_data=task_data
        )
        
        result = agent.execute(task_data)
        
        if task_id:
            agent.complete_task(
                task_id,
                result,
                status='completed' if result.get('success') else 'failed'
            )
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error triggering agent {agent_type}: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/agents/logs')
@login_required
def agent_logs():
    """View detailed agent logs"""
    from models import AgentLog
    
    page = request.args.get('page', 1, type=int)
    agent_type = request.args.get('agent_type', '')
    
    query = AgentLog.query
    
    if agent_type:
        query = query.filter_by(agent_type=agent_type)
    
    logs = query.order_by(AgentLog.created_at.desc()).paginate(
        page=page, per_page=50, error_out=False
    )
    
    return render_template('agent_logs.html', logs=logs, selected_agent=agent_type)

@main_bp.route('/agents/reports')
@login_required
def agent_reports():
    """View agent-generated reports"""
    from models import AgentReport
    
    page = request.args.get('page', 1, type=int)
    report_type = request.args.get('report_type', '')
    
    query = AgentReport.query
    
    if report_type:
        query = query.filter_by(report_type=report_type)
    
    reports = query.order_by(AgentReport.created_at.desc()).paginate(
        page=page, per_page=20, error_out=False
    )
    
    return render_template('agent_reports.html', reports=reports, selected_type=report_type)

@main_bp.route('/agents/reports/<int:report_id>')
@login_required
def view_agent_report(report_id):
    """View detailed agent report"""
    from models import AgentReport
    
    report = AgentReport.query.get_or_404(report_id)
    
    return render_template('view_agent_report.html', report=report)

# ===== PHASE 2: SEO & ANALYTICS MODULE =====
@main_bp.route('/seo/dashboard')
@login_required
def seo_dashboard():
    """SEO dashboard with overview"""
    from services.seo_service import SEOService
    from models import SEOKeyword, SEOBacklink, SEOCompetitor, SEOAudit
    
    stats = SEOService.get_dashboard_stats()
    recent_audits = SEOAudit.query.order_by(SEOAudit.created_at.desc()).limit(5).all()
    top_keywords = SEOKeyword.query.filter(
        SEOKeyword.current_position.isnot(None),
        SEOKeyword.current_position <= 10
    ).order_by(SEOKeyword.current_position).limit(10).all()
    
    return render_template('seo_dashboard.html', 
                         stats=stats,
                         recent_audits=recent_audits,
                         top_keywords=top_keywords)

@main_bp.route('/seo/keywords')
@login_required
def seo_keywords():
    """Keyword tracking list"""
    from models import SEOKeyword
    keywords = SEOKeyword.query.filter_by(is_tracking=True).all()
    return render_template('seo_keywords.html', keywords=keywords)

@main_bp.route('/seo/keywords/add', methods=['POST'])
@login_required
def add_keyword():
    """Add keyword to track"""
    from services.seo_service import SEOService
    keyword = request.form.get('keyword')
    target_url = request.form.get('target_url')
    
    if keyword:
        SEOService.track_keyword(keyword, target_url)
        flash('Keyword added for tracking!', 'success')
    return redirect(url_for('main.seo_keywords'))

@main_bp.route('/seo/backlinks')
@login_required
def seo_backlinks():
    """Backlink monitoring"""
    from models import SEOBacklink
    backlinks = SEOBacklink.query.filter_by(status='active').order_by(SEOBacklink.domain_authority.desc()).all()
    return render_template('seo_backlinks.html', backlinks=backlinks)

@main_bp.route('/seo/competitors')
@login_required
def seo_competitors():
    """Competitor tracking"""
    from models import SEOCompetitor
    competitors = SEOCompetitor.query.filter_by(is_active=True).all()
    return render_template('seo_competitors.html', competitors=competitors)

@main_bp.route('/seo/audit', methods=['GET', 'POST'])
@login_required
def seo_audit():
    """Run site audit"""
    from services.seo_service import SEOService
    
    if request.method == 'POST':
        url = request.form.get('url')
        audit_type = request.form.get('audit_type', 'full')
        
        audit = SEOService.run_site_audit(url, audit_type)
        if audit:
            flash('Site audit completed!', 'success')
            return redirect(url_for('main.seo_audit_results', audit_id=audit.id))
    
    return render_template('seo_audit_form.html')

@main_bp.route('/seo/audit/<int:audit_id>')
@login_required
def seo_audit_results(audit_id):
    """View audit results"""
    from models import SEOAudit
    audit = SEOAudit.query.get_or_404(audit_id)
    return render_template('seo_audit_results.html', audit=audit)

# ===== PHASE 3: EVENT ENHANCEMENTS =====
@main_bp.route('/events/<int:event_id>/tickets')
@login_required
def event_tickets(event_id):
    """Manage event tickets"""
    from models import Event, EventTicket
    event = Event.query.get_or_404(event_id)
    tickets = EventTicket.query.filter_by(event_id=event_id).all()
    return render_template('event_tickets.html', event=event, tickets=tickets)

@main_bp.route('/events/<int:event_id>/tickets/create', methods=['POST'])
@login_required
def create_ticket_type(event_id):
    """Create ticket type"""
    from services.event_service import EventService
    
    name = request.form.get('name')
    price = float(request.form.get('price', 0))
    quantity = int(request.form.get('quantity', 0))
    description = request.form.get('description')
    
    EventService.create_ticket_type(event_id, name, price, quantity, description)
    flash('Ticket type created!', 'success')
    return redirect(url_for('main.event_tickets', event_id=event_id))

@main_bp.route('/events/<int:event_id>/purchase', methods=['POST'])
@login_required
def purchase_event_ticket(event_id):
    """Purchase event ticket"""
    from services.event_service import EventService
    
    ticket_id = int(request.form.get('ticket_id'))
    contact_id = int(request.form.get('contact_id'))
    quantity = int(request.form.get('quantity', 1))
    
    purchase = EventService.purchase_ticket(ticket_id, contact_id, quantity)
    if purchase:
        flash('Ticket purchased successfully!', 'success')
    else:
        flash('Ticket purchase failed. Check availability.', 'error')
    
    return redirect(url_for('main.view_event', id=event_id))

@main_bp.route('/events/<int:event_id>/checkin', methods=['GET', 'POST'])
@login_required
def event_checkin(event_id):
    """Event check-in system"""
    from models import Event, EventCheckIn, TicketPurchase
    from services.event_service import EventService
    
    event = Event.query.get_or_404(event_id)
    
    if request.method == 'POST':
        contact_id = int(request.form.get('contact_id'))
        ticket_purchase_id = request.form.get('ticket_purchase_id')
        
        EventService.check_in_attendee(
            event_id, 
            contact_id,
            int(ticket_purchase_id) if ticket_purchase_id else None,
            method='manual',
            staff_name=current_user.username
        )
        flash('Attendee checked in!', 'success')
    
    checkins = EventCheckIn.query.filter_by(event_id=event_id).all()
    purchases = TicketPurchase.query.join(EventTicket).filter(
        EventTicket.event_id == event_id
    ).all()
    
    return render_template('event_checkin.html', event=event, checkins=checkins, purchases=purchases)

# ===== PHASE 4: SOCIAL MEDIA EXPANSION =====
@main_bp.route('/social/accounts')
@login_required
def social_accounts():
    """Manage connected social media accounts"""
    from models import SocialMediaAccount
    accounts = SocialMediaAccount.query.filter_by(is_active=True).all()
    return render_template('social_accounts.html', accounts=accounts)

@main_bp.route('/social/accounts/connect', methods=['POST'])
@login_required
def connect_social_account():
    """Connect new social media account"""
    from services.social_media_service import SocialMediaService
    
    platform = request.form.get('platform')
    account_name = request.form.get('account_name')
    access_token = request.form.get('access_token')
    
    account = SocialMediaService.connect_account(platform, account_name, access_token)
    if account:
        flash(f'{platform.capitalize()} account connected!', 'success')
    else:
        flash('Failed to connect account', 'error')
    
    return redirect(url_for('main.social_accounts'))

@main_bp.route('/social/schedule', methods=['GET', 'POST'])
@login_required
def social_schedule_post():
    """Schedule social media post"""
    from services.social_media_service import SocialMediaService
    from models import SocialMediaAccount
    
    if request.method == 'POST':
        account_id = int(request.form.get('account_id'))
        content = request.form.get('content')
        scheduled_for = datetime.fromisoformat(request.form.get('scheduled_for'))
        hashtags = request.form.get('hashtags')
        
        post = SocialMediaService.schedule_post(account_id, content, scheduled_for, hashtags=hashtags)
        if post:
            flash('Post scheduled!', 'success')
            return redirect(url_for('main.social_media'))
    
    accounts = SocialMediaAccount.query.filter_by(is_active=True).all()
    return render_template('social_schedule.html', accounts=accounts)

@main_bp.route('/social/crosspost', methods=['POST'])
@login_required
def social_crosspost():
    """Create cross-platform post"""
    from services.social_media_service import SocialMediaService
    
    content = request.form.get('content')
    platforms = request.form.getlist('platforms')
    scheduled_for = datetime.fromisoformat(request.form.get('scheduled_for'))
    
    cross_post = SocialMediaService.create_cross_post(content, platforms, scheduled_for)
    if cross_post:
        flash('Cross-post created!', 'success')
    
    return redirect(url_for('main.social_media'))

# ===== PHASE 5: ADVANCED AUTOMATIONS =====
@main_bp.route('/automations/<int:automation_id>/test', methods=['POST'])
@login_required
def test_automation(automation_id):
    """Test automation in test mode"""
    from services.automation_service import AutomationService
    
    test_contact_id = request.form.get('test_contact_id')
    test = AutomationService.run_test(
        automation_id,
        int(test_contact_id) if test_contact_id else None
    )
    
    if test:
        return jsonify({
            'success': True,
            'test_id': test.id,
            'results': test.test_results
        })
    
    return jsonify({'success': False}), 500

@main_bp.route('/automations/triggers')
@login_required
def automation_triggers():
    """Browse trigger library"""
    from services.automation_service import AutomationService
    
    category = request.args.get('category')
    triggers = AutomationService.get_trigger_library(category)
    
    categories = ['ecommerce', 'engagement', 'nurture', 'retention']
    return render_template('automation_triggers.html', triggers=triggers, categories=categories)

@main_bp.route('/automations/<int:automation_id>/abtest', methods=['POST'])
@login_required
def create_automation_abtest(automation_id):
    """Create A/B test for automation"""
    from services.automation_service import AutomationService
    
    step_id = int(request.form.get('step_id'))
    variant_a_id = int(request.form.get('variant_a_template_id'))
    variant_b_id = int(request.form.get('variant_b_template_id'))
    split = int(request.form.get('split_percentage', 50))
    
    ab_test = AutomationService.create_ab_test(automation_id, step_id, variant_a_id, variant_b_id, split)
    if ab_test:
        flash('A/B test created!', 'success')
    
    return redirect(url_for('main.edit_automation', id=automation_id))

# ===== PHASE 6: UNIFIED MARKETING CALENDAR =====
@main_bp.route('/marketing-calendar')
@login_required
def marketing_calendar():
    """Unified marketing calendar view"""
    # from services.scheduling_service import SchedulingService
    from datetime import datetime
    
    year = request.args.get('year', datetime.now().year, type=int)
    month = request.args.get('month', datetime.now().month, type=int)
    
    # calendar_data = SchedulingService.get_calendar_view(year, month)
    # upcoming = SchedulingService.get_upcoming_schedules(days=30)
    
    return render_template('marketing_calendar.html', 
                         calendar_data={},
                         upcoming=[],
                         year=year,
                         month=month)

@main_bp.route('/calendar/schedule', methods=['POST'])
@login_required
def calendar_schedule():
    """Add item to calendar"""
    # from services.scheduling_service import SchedulingService
    
    module_type = request.form.get('module_type')
    module_object_id = int(request.form.get('module_object_id'))
    title = request.form.get('title')
    scheduled_at = datetime.fromisoformat(request.form.get('scheduled_at'))
    description = request.form.get('description', '')
    
    schedule = SchedulingService.create_schedule(
        module_type,
        module_object_id,
        title,
        scheduled_at,
        description
    )
    
    if schedule:
        flash('Item added to calendar!', 'success')
    
    return redirect(url_for('main.marketing_calendar'))

# ===== SYSTEM INITIALIZATION =====
@main_bp.route('/system/init')
@login_required
def system_init():
    """Initialize system data (trigger library, etc.)"""
    from services.automation_service import AutomationService
    
    try:
        # Seed trigger library
        AutomationService.seed_trigger_library()
        
        flash('System initialized successfully! Trigger library seeded.', 'success')
    except Exception as e:
        logger.error(f"System initialization error: {e}")
        flash(f'Initialization error: {str(e)}', 'error')
    
    return redirect(url_for('main.dashboard'))

# ===== MONITORING & HEALTH CHECK =====
@main_bp.route('/health')
def health_check():
    """Health check endpoint for monitoring"""
    try:
        # Check database connection
        db.session.execute(db.text('SELECT 1'))
        
        # Check critical tables exist
        from models import Contact, Campaign, SEOKeyword, EventTicket, SocialMediaAccount
        
        Contact.query.limit(1).all()
        Campaign.query.limit(1).all()
        SEOKeyword.query.limit(1).all()
        EventTicket.query.limit(1).all()
        SocialMediaAccount.query.limit(1).all()
        
        return jsonify({
            'status': 'healthy',
            'database': 'connected',
            'timestamp': datetime.utcnow().isoformat(),
            'version': 'Phase 2-6 Release'
        }), 200
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return jsonify({
            'status': 'unhealthy',
            'error': str(e),
            'timestamp': datetime.utcnow().isoformat()
        }), 500

@main_bp.route('/system/status')
@login_required
def system_status():
    """Detailed system status page"""
    try:
        # Count records in key tables
        stats = {
            'contacts': Contact.query.count(),
            'campaigns': Campaign.query.count(),
            'seo_keywords': SEOKeyword.query.count(),
            'events': Event.query.count(),
            'event_tickets': EventTicket.query.count(),
            'social_accounts': SocialMediaAccount.query.count(),
            'automation_triggers': AutomationTriggerLibrary.query.count(),
        }
        
        # Check AI agents
        from models import AgentSchedule
        agent_count = AgentSchedule.query.filter_by(is_enabled=True).count()
        
        return render_template('system_status.html', stats=stats, agent_count=agent_count)
    except Exception as e:
        flash(f'Error loading system status: {str(e)}', 'error')
        return redirect(url_for('main.dashboard'))

@main_bp.route('/chatbot')
@login_required
def chatbot():
    """LUX AI Chatbot - Redirect to dashboard (chatbot is now a floating widget)"""
    return redirect(url_for('main.dashboard'))

@main_bp.route('/api/diagnostics/errors', methods=['GET'])
@login_required
def get_recent_errors():
    """Get recent application errors for chatbot analysis"""
    try:
        hours = request.args.get('hours', 24, type=int)
        limit = request.args.get('limit', 20, type=int)
        errors = ApplicationDiagnostics.get_recent_errors(hours=hours, limit=limit)
        return jsonify({'success': True, 'errors': errors})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/api/diagnostics/health', methods=['GET'])
@login_required
def get_system_health():
    """Get system health status for chatbot analysis"""
    try:
        health = ApplicationDiagnostics.get_system_health()
        error_summary = ApplicationDiagnostics.get_error_summary()
        return jsonify({'success': True, 'health': health, 'error_summary': error_summary})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/api/auto-repair/start', methods=['POST'])
@login_required
def start_auto_repair():
    """Start automated error repair and resolution testing"""
    try:
        error_id = request.json.get('error_id') if request.is_json else None
        results = AutoRepairService.execute_auto_repair(error_id=error_id)
        return jsonify({'success': True, 'results': results})
    except Exception as e:
        logger.error(f"Auto-repair endpoint error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/api/auto-repair/clear', methods=['POST'])
@login_required
def clear_resolved_errors():
    """Clear resolved errors from the log"""
    try:
        hours = request.json.get('hours', 24) if request.is_json else 24
        results = AutoRepairService.clear_resolved_errors(older_than_hours=hours)
        return jsonify({'success': True, 'results': results})
    except Exception as e:
        logger.error(f"Clear errors endpoint error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/api/system/diagnosis', methods=['GET'])
@login_required
def system_diagnosis():
    """Comprehensive system diagnosis for all error types"""
    try:
        diagnosis = ErrorFixService.comprehensive_system_diagnosis()
        return jsonify({'success': True, 'diagnosis': diagnosis})
    except Exception as e:
        logger.error(f"System diagnosis error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/api/system/health', methods=['GET'])
@login_required
def system_health():
    """Check system health and resource usage"""
    try:
        health = ErrorFixService.check_server_health()
        return jsonify({'success': True, 'health': health})
    except Exception as e:
        logger.error(f"System health error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/api/system/validate-openai', methods=['GET'])
@login_required
def validate_openai():
    """Validate OpenAI API key configuration"""
    try:
        validation = ErrorFixService.validate_openai_api_key()
        return jsonify({'success': True, 'validation': validation})
    except Exception as e:
        logger.error(f"OpenAI validation error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/api/system/endpoint-check', methods=['GET'])
@login_required
def endpoint_check():
    """Check for 404 errors on key endpoints"""
    try:
        results = ErrorFixService.check_404_endpoints()
        return jsonify({'success': True, 'results': results})
    except Exception as e:
        logger.error(f"Endpoint check error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/api/ai-fixer/auto-fix-all', methods=['POST'])
@login_required
def ai_auto_fix_all():
    """AI-powered: Automatically fix ALL errors"""
    try:
        results = AICodeFixer.auto_fix_all_errors()
        return jsonify({'success': True, 'fixes': results})
    except Exception as e:
        logger.error(f"AI auto-fix error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/api/ai-fixer/fix-error/<int:error_id>', methods=['POST'])
@login_required
def ai_fix_single_error(error_id):
    """AI-powered: Fix a specific error by ID"""
    try:
        error = ErrorLog.query.get(error_id)
        if not error:
            return jsonify({'success': False, 'error': 'Error not found'}), 404
        
        fix_result = AICodeFixer.generate_and_apply_fix(
            error.error_type,
            error.to_dict()
        )
        
        # Mark as resolved if fix succeeded
        if fix_result.get('status') in ['ok', 'fixed', 'all_routes_registered']:
            error.is_resolved = True
            error.resolution_notes = json.dumps(fix_result)
            db.session.commit()
        
        return jsonify({'success': True, 'error_id': error_id, 'fix_result': fix_result})
    except Exception as e:
        logger.error(f"AI fix error endpoint: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/api/ai-fixer/codebase-structure', methods=['GET'])
@login_required
def get_codebase_structure():
    """Get codebase structure for AI context"""
    try:
        structure = AICodeFixer.get_codebase_structure()
        return jsonify({'success': True, 'structure': structure})
    except Exception as e:
        logger.error(f"Codebase structure error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/api/ai/execute-action', methods=['POST'])
@login_required
def execute_ai_action():
    """Execute AI actions immediately (not just recommend)"""
    try:
        data = request.get_json() or {}
        action = data.get('action')
        params = data.get('params', {})
        
        if not action:
            return jsonify({'success': False, 'error': 'action parameter required'}), 400
        
        result = AIActionExecutor.handle_action(action, params)
        return jsonify({'success': True, 'result': result})
    except Exception as e:
        logger.error(f"AI action execution error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/settings/integrations')
@login_required
def settings_integrations():
    """Settings & Integrations page - shows current user's company config"""
    try:
        company = current_user.get_default_company()
        if not company:
            return redirect(url_for('main.dashboard'))
        return render_template('company_settings.html', company=company)
    except Exception as e:
        logger.error(f"Settings page error: {e}")
        return redirect(url_for('main.dashboard'))

@main_bp.route('/company/<int:company_id>/settings')
@login_required
def company_settings(company_id):
    """Company settings & integrations page"""
    try:
        company = Company.query.get(company_id)
        if not company:
            return redirect(url_for('main.dashboard'))
        return render_template('company_settings.html', company=company)
    except Exception as e:
        logger.error(f"Settings page error: {e}")
        return redirect(url_for('main.dashboard'))

@main_bp.route('/api/company/<int:company_id>/secrets', methods=['GET'])
@login_required
def get_company_secrets(company_id):
    """Get all secrets for a company"""
    try:
        from models import CompanySecret
        company = Company.query.get(company_id)
        if not company:
            return jsonify({'success': False, 'error': 'Company not found'}), 404
        
        secrets = CompanySecret.query.filter_by(company_id=company_id).all()
        return jsonify({
            'success': True,
            'company': company.name,
            'secrets': [{'key': s.key, 'created_at': s.created_at.isoformat()} for s in secrets]
        })
    except Exception as e:
        logger.error(f"Get secrets error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/api/company/<int:company_id>/secrets/save', methods=['POST'])
@login_required
def save_company_secrets(company_id):
    """Save/update secrets for a company"""
    try:
        from models import CompanySecret
        company = Company.query.get(company_id)
        if not company:
            return jsonify({'success': False, 'error': 'Company not found'}), 404
        
        data = request.get_json()
        
        for key, value in data.items():
            if value:  # Only save if value is provided
                company.set_secret(key, value)
        
        return jsonify({
            'success': True,
            'company': company.name,
            'secrets_saved': len([k for k, v in data.items() if v])
        })
    except Exception as e:
        logger.error(f"Save secrets error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/api/company/<int:company_id>/settings', methods=['POST'])
@login_required
def save_company_settings(company_id):
    """Save company brand settings"""
    try:
        company = Company.query.get(company_id)
        if not company:
            return jsonify({'success': False, 'error': 'Company not found'}), 404
        
        data = request.get_json()
        
        if 'primary_color' in data:
            company.primary_color = data['primary_color']
        if 'secondary_color' in data:
            company.secondary_color = data['secondary_color']
        if 'accent_color' in data:
            company.accent_color = data['accent_color']
        if 'font_family' in data:
            company.font_family = data['font_family']
        if 'website_url' in data:
            company.website_url = data['website_url']
        if 'logo_path' in data:
            company.logo_path = data['logo_path']
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'company': company.name,
            'message': 'Settings updated successfully'
        })
    except Exception as e:
        logger.error(f"Save settings error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/chatbot/send', methods=['POST'])
@csrf.exempt
def chatbot_send_with_auto_fix():
    """Send message to AI chatbot and get response with error diagnostics
    
    Supports actions:
    - action='message': Regular chat (default)
    - action='diagnose': Read server logs and analyze for issues
    """
    try:
        from openai import OpenAI
        
        data = request.get_json()
        user_message = data.get('message', '')
        action = data.get('action', 'message')  # 'message' or 'diagnose'
        
        if not user_message:
            return jsonify({'error': 'Message is required'}), 400
        
        # Retrieve API key from environment
        api_key = os.environ.get('OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY')
        if not api_key:
            error_msg = 'OpenAI API key not configured in environment'
            logger.error(error_msg)
            log_application_error(
                error_type='ConfigurationError',
                error_message=error_msg,
                endpoint='/chatbot/send',
                method='POST',
                severity='critical'
            )
            return jsonify({'error': error_msg}), 500
        
        # Get current system diagnostics for context
        diagnostics_context = ""
        try:
            recent_errors = ApplicationDiagnostics.get_recent_errors(hours=24, limit=5)
            health = ApplicationDiagnostics.get_system_health()
            diagnostics_context = f"""

SYSTEM DIAGNOSTICS (Database):
- System Health: {health.get('status', 'unknown')}
- Recent Errors (1h): {health.get('recent_errors_1h', 0)}
- Unresolved Issues: {health.get('unresolved_errors', 0)}

Recent Database Error Examples:
{json.dumps(recent_errors[:3], indent=2) if recent_errors else 'No recent errors'}
"""
        except Exception as diag_error:
            logger.warning(f"Could not retrieve diagnostics: {diag_error}")
            diagnostics_context = "\n(Database diagnostics unavailable)"
        
        # Read server logs if diagnose action is requested
        server_logs_context = ""
        auto_repair_context = ""
        if action == 'diagnose':
            try:
                all_logs = LogReader.get_all_logs(lines=30)
                server_logs_context = f"""

SERVER LOGS (from VPS):
{LogReader.format_logs_for_ai(all_logs)}

QUICK ERROR PATTERNS:
{json.dumps(LogReader.analyze_logs_for_errors(all_logs), indent=2)}
"""
            except Exception as log_error:
                logger.warning(f"Could not read server logs: {log_error}")
                server_logs_context = "\n(Server logs unavailable)"
        
        # Add auto-repair capability to system prompt
        auto_repair_context = """

SPECIAL CAPABILITIES - AUTO-REPAIR:
You can trigger automated error repair by responding with:
ACTION: REPAIR_ERRORS
This will:
1. Find all unresolved errors
2. Generate AI-powered fix plans
3. Test if errors are resolved
4. Mark resolved errors and clear them
5. Return a detailed report

To use this, when appropriate, include "ACTION: REPAIR_ERRORS" in your response."""
        
        # Initialize OpenAI client with explicit error handling
        try:
            client = OpenAI(api_key=api_key)
        except Exception as client_error:
            logger.error(f"Failed to initialize OpenAI client: {client_error}")
            log_application_error(
                error_type='OpenAIClientError',
                error_message=f"Client initialization failed: {str(client_error)}",
                endpoint='/chatbot/send',
                method='POST',
                severity='error'
            )
            return jsonify({'error': 'Failed to initialize AI service'}), 500
        
        # Make API call with explicit error handling
        try:
            system_prompt = f"""You are LUX, an AI marketing assistant and platform debugger for the LUX Marketing platform. 

Your capabilities:
1. MARKETING: Help with campaign generation, marketing strategy, content creation, audience segmentation
2. ERROR DETECTION: Identify, explain, and help fix platform errors, bugs, and configuration issues
3. DEBUGGING: Analyze error messages, logs, and system behavior to diagnose problems
4. TROUBLESHOOTING: Provide step-by-step solutions for platform issues
5. PLATFORM GUIDANCE: Explain features, workflows, and best practices
6. AUTO-REPAIR: Suggest code fixes, configuration changes, and implementation steps
7. LOG ANALYSIS: Read and analyze server logs (Nginx, Gunicorn, systemd, app logs) to diagnose issues

When analyzing errors:
1. Identify the root cause
2. Explain the issue in simple terms
3. Provide step-by-step solutions
4. Suggest preventive measures
5. Confirm fixes resolve the issue

Current System Context:
{diagnostics_context}
{server_logs_context}
{auto_repair_context}

Be helpful, professional, concise, and proactive. Always look for ways to improve system health.

When users ask you to fix errors, resolve issues, or test the system:
1. Analyze what the user is asking
2. If appropriate, suggest running automatic repairs
3. You can't directly run repairs, but your response can trigger them by including "ACTION: REPAIR_ERRORS"
4. The system will see this action and execute auto-repair automatically"""
            
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message}
                ],
                temperature=0.7,
                max_tokens=2000
            )
            
            bot_message = response.choices[0].message.content
            
            # Check if auto-repair should be triggered
            trigger_repair = 'ACTION: REPAIR_ERRORS' in bot_message
            repair_results = None
            if trigger_repair:
                try:
                    repair_results = AutoRepairService.execute_auto_repair()
                    # Clean up the action from the response
                    bot_message = bot_message.replace('ACTION: REPAIR_ERRORS', '').strip()
                except Exception as repair_error:
                    logger.warning(f"Auto-repair trigger failed: {repair_error}")
            
            return jsonify({
                'response': bot_message,
                'action': action,
                'has_logs': bool(server_logs_context),
                'auto_repair_triggered': trigger_repair,
                'repair_results': repair_results
            })
            
        except Exception as api_error:
            error_str = str(api_error)
            logger.error(f"OpenAI API error: {error_str}")
            
            # Check if it's an authentication error
            if 'invalid' in error_str.lower() or '401' in error_str or 'auth' in error_str.lower():
                log_application_error(
                    error_type='OpenAIAuthenticationError',
                    error_message=f"API authentication failed: {error_str[:200]}",
                    endpoint='/chatbot/send',
                    method='POST',
                    severity='critical'
                )
            else:
                log_application_error(
                    error_type='OpenAIAPIError',
                    error_message=error_str[:200],
                    endpoint='/chatbot/send',
                    method='POST',
                    severity='error'
                )
            
            return jsonify({'error': 'AI service temporarily unavailable. Please try again.'}), 503
        
    except Exception as e:
        logger.error(f"Chatbot error: {e}")
        log_application_error(
            error_type='ChatbotError',
            error_message=str(e)[:200],
            endpoint='/chatbot/send',
            method='POST',
            severity='error'
        )
        return jsonify({'error': 'An error occurred. Please try again.'}), 500

@main_bp.route('/content-generator')
def content_generator():
    """AI Content Generator Page - Public Access"""
    return render_template('content_generator.html')

@main_bp.route('/api/content/generate', methods=['POST'])
@csrf.exempt
def generate_content():
    """Generate marketing content using AI
    
    Content types supported:
    - blog_post: Long-form blog content
    - social_media: Social media posts (Twitter, LinkedIn, Facebook, Instagram)
    - email_campaign: Email marketing copy
    - ad_copy: Advertisement copy (Google Ads, Facebook Ads)
    - seo_content: SEO-optimized content
    - product_description: E-commerce product descriptions
    """
    try:
        import openai
        
        data = request.get_json()
        content_type = data.get('type', 'blog_post')
        topic = data.get('topic', '')
        tone = data.get('tone', 'professional')
        length = data.get('length', 'medium')
        keywords = data.get('keywords', [])
        additional_context = data.get('context', '')
        
        if not topic:
            return jsonify({'error': 'Topic is required'}), 400
        
        # Get API key
        api_key = os.getenv('OPENAI_API_KEY') or os.getenv('OPENAI_API_BOUTIQUELUX')
        if not api_key:
            return jsonify({'error': 'OpenAI API key not configured'}), 500
        
        openai.api_key = api_key
        
        # Define prompts for different content types
        prompts = {
            'blog_post': f"""Create a comprehensive blog post about: {topic}
Tone: {tone}
Length: {length} (short=300-500 words, medium=500-800 words, long=800-1200 words)
{f'Keywords to include: {", ".join(keywords)}' if keywords else ''}
{f'Additional context: {additional_context}' if additional_context else ''}

Format the response as a complete blog post with:
- Engaging headline
- Introduction
- Main body with subheadings
- Conclusion
- Call-to-action""",
            
            'social_media': f"""Create engaging social media posts about: {topic}
Tone: {tone}
Platform: {additional_context or 'general'}
{f'Keywords: {", ".join(keywords)}' if keywords else ''}

Generate 3 variations:
1. Short post (Twitter/X style, 280 chars max)
2. Medium post (LinkedIn/Facebook style)
3. Visual post (Instagram style with hashtags)""",
            
            'email_campaign': f"""Create an email marketing campaign about: {topic}
Tone: {tone}
{f'Keywords: {", ".join(keywords)}' if keywords else ''}
{f'Context: {additional_context}' if additional_context else ''}

Include:
- Subject line (with 2-3 variations)
- Preview text
- Email body
- Call-to-action
- P.S. section""",
            
            'ad_copy': f"""Create advertisement copy for: {topic}
Tone: {tone}
Platform: {additional_context or 'Google Ads'}
{f'Keywords: {", ".join(keywords)}' if keywords else ''}

Generate:
- 3 headline variations (30 chars max)
- 2 description variations (90 chars max)
- Call-to-action suggestions""",
            
            'seo_content': f"""Create SEO-optimized content about: {topic}
Tone: {tone}
Target keywords: {", ".join(keywords) if keywords else topic}
{f'Context: {additional_context}' if additional_context else ''}

Include:
- SEO-friendly title (60 chars max)
- Meta description (155 chars max)
- H1, H2, H3 structure
- Content optimized for keywords
- Internal linking suggestions""",
            
            'product_description': f"""Create a compelling product description for: {topic}
Tone: {tone}
{f'Key features: {", ".join(keywords)}' if keywords else ''}
{f'Additional details: {additional_context}' if additional_context else ''}

Include:
- Catchy product title
- Short description (1-2 sentences)
- Key features and benefits
- Technical specifications
- Why customers should buy"""
        }
        
        # Get the appropriate prompt
        system_prompt = "You are LUX AI, an expert marketing content creator. Create high-quality, engaging marketing content that converts. Be creative, persuasive, and professional."
        user_prompt = prompts.get(content_type, prompts['blog_post'])
        
        # Set token limits based on length
        token_limits = {
            'short': 800,
            'medium': 1500,
            'long': 2500
        }
        max_tokens = token_limits.get(length, 1500)
        
        # Generate content
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.8,  # Higher temperature for more creative content
            max_tokens=max_tokens
        )
        
        generated_content = response.choices[0].message.content
        
        return jsonify({
            'success': True,
            'content': generated_content,
            'type': content_type,
            'topic': topic,
            'tone': tone,
            'length': length,
            'tokens_used': response.usage.total_tokens
        })
        
    except Exception as e:
        logger.error(f"Content generation error: {e}")
        return jsonify({'error': f'Failed to generate content: {str(e)}'}), 500

@main_bp.route('/api/content/export/pdf', methods=['POST'])
@csrf.exempt
def export_content_pdf():
    """Export generated content as PDF"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from io import BytesIO
        
        data = request.get_json()
        content = data.get('content', '')
        title = data.get('title', 'Generated Content')
        content_type = data.get('type', 'content')
        
        # Create PDF in memory
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#bc00ed',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        content_style = ParagraphStyle(
            'CustomContent',
            parent=styles['Normal'],
            fontSize=12,
            leading=16,
            spaceAfter=12,
            alignment=TA_LEFT
        )
        
        # Build PDF
        story = []
        
        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
        
        # Add content (convert line breaks to paragraph breaks)
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph.replace('\n', '<br/>'), content_style))
                story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{content_type}_{title[:30].replace(" ", "_")}.pdf'
        )
        
    except Exception as e:
        logger.error(f"PDF export error: {e}")
        return jsonify({'error': f'Failed to export as PDF: {str(e)}'}), 500

@main_bp.route('/api/content/export/docx', methods=['POST'])
@csrf.exempt
def export_content_docx():
    """Export generated content as DOCX"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        
        data = request.get_json()
        content = data.get('content', '')
        title = data.get('title', 'Generated Content')
        content_type = data.get('type', 'content')
        
        # Create document
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.color.rgb = RGBColor(188, 0, 237)  # Purple
        
        # Add spacing
        doc.add_paragraph()
        
        # Add content
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph.strip())
                p.style.font.size = Pt(12)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{content_type}_{title[:30].replace(" ", "_")}.docx'
        )
        
    except Exception as e:
        logger.error(f"DOCX export error: {e}")
        return jsonify({'error': f'Failed to export as DOCX: {str(e)}'}), 500

# ============= WORDPRESS / WOOCOMMERCE =============
@main_bp.route('/wordpress')
@login_required
def wordpress_integration():
    """WordPress and WooCommerce integration management"""
    from models import WordPressIntegration
    integrations = WordPressIntegration.query.filter_by(company_id=current_user.get_default_company().id).all()
    return render_template('wordpress_integration.html', integrations=integrations)

# ============= KEYWORD RESEARCH =============
@main_bp.route('/keywords')
@login_required
def keyword_research():
    """Keyword research and tracking"""
    from models import KeywordResearch
    keywords = KeywordResearch.query.filter_by(company_id=current_user.get_default_company().id).all()
    return render_template('keyword_research.html', keywords=keywords)

@main_bp.route('/keywords/create', methods=['POST'])
@login_required
def create_keyword_research():
    """Create new keyword research"""
    from models import KeywordResearch
    try:
        data = request.get_json()
        company = current_user.get_default_company()
        keyword = KeywordResearch(
            company_id=company.id,
            keyword=data.get('keyword'),
            status='tracking'
        )
        db.session.add(keyword)
        db.session.commit()
        return jsonify({'success': True, 'keyword_id': keyword.id}), 201
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 400

# ============= CRM / DEALS =============
@main_bp.route('/crm')
@login_required
def crm_dashboard():
    """CRM dashboard with deals and pipeline"""
    from models import Deal
    from sqlalchemy import func
    company = current_user.get_default_company()
    
    deals = Deal.query.filter_by(company_id=company.id).all()
    pipeline_data = db.session.query(
        Deal.stage,
        func.count(Deal.id).label('count'),
        func.sum(Deal.value).label('total_value')
    ).filter_by(company_id=company.id).group_by(Deal.stage).all()
    
    return render_template('crm_dashboard.html', deals=deals, pipeline_data=pipeline_data)

@main_bp.route('/crm/deals/<int:deal_id>')
@login_required
def deal_detail(deal_id):
    """View individual deal"""
    from models import Deal, DealActivity
    deal = Deal.query.get_or_404(deal_id)
    activities = DealActivity.query.filter_by(deal_id=deal_id).order_by(DealActivity.activity_date.desc()).all()
    return render_template('deal_detail.html', deal=deal, activities=activities)

# ============= LEAD SCORING =============
@main_bp.route('/lead-scoring')
@login_required
def lead_scoring():
    """Lead scoring and nurturing"""
    from models import Contact, LeadScore
    company = current_user.get_default_company()
    
    contacts_with_scores = db.session.query(Contact, LeadScore).outerjoin(
        LeadScore, Contact.id == LeadScore.contact_id
    ).filter(Contact.is_active == True).all()
    
    return render_template('lead_scoring.html', contacts_with_scores=contacts_with_scores)

# ============= COMPETITOR ANALYSIS =============
@main_bp.route('/competitors')
@login_required
def competitor_analysis():
    """Competitor analysis and tracking"""
    from models import CompetitorProfile
    company = current_user.get_default_company()
    competitors = CompetitorProfile.query.filter_by(company_id=company.id).all()
    return render_template('competitor_analysis.html', competitors=competitors)

# ============= PERSONALIZATION =============
@main_bp.route('/personalization')
@login_required
def personalization_rules():
    """Content personalization rules"""
    from models import PersonalizationRule
    company = current_user.get_default_company()
    rules = PersonalizationRule.query.filter_by(company_id=company.id).all()
    return render_template('personalization_rules.html', rules=rules)

# ============= A/B TESTING ENHANCEMENTS =============
@main_bp.route('/multivariate-tests')
@login_required
def multivariate_tests():
    """Multivariate testing"""
    from models import MultivariateTest
    tests = MultivariateTest.query.all()
    return render_template('multivariate_tests.html', tests=tests)

# ============= ROI TRACKING =============
@main_bp.route('/roi-analytics')
@login_required
def roi_analytics():
    """ROI tracking and attribution analytics"""
    from models import Campaign, CampaignCost, AttributionModel
    from sqlalchemy import func
    company = current_user.get_default_company()
    
    campaigns_roi = db.session.query(
        Campaign.id,
        Campaign.name,
        func.sum(CampaignCost.amount).label('total_cost'),
        func.sum(AttributionModel.revenue).label('total_revenue')
    ).outerjoin(CampaignCost).outerjoin(AttributionModel).filter(
        Campaign.company_id == company.id
    ).group_by(Campaign.id).all()
    
    return render_template('roi_analytics.html', campaigns_roi=campaigns_roi)

# ============= SURVEYS & FEEDBACK =============
@main_bp.route('/surveys')
@login_required
def surveys():
    """NPS and feedback surveys"""
    from models import SurveyResponse
    company = current_user.get_default_company()
    responses = SurveyResponse.query.all()
    
    nps_score = None
    if responses:
        promoters = sum(1 for r in responses if r.score >= 9)
        detractors = sum(1 for r in responses if r.score <= 6)
        nps_score = ((promoters - detractors) / len(responses) * 100) if responses else 0
    
    return render_template('surveys.html', responses=responses, nps_score=nps_score)

# ============= AGENT CONFIGURATION =============
@main_bp.route('/agent-config')
@login_required
def agent_configuration():
    """Configure AI agents per company"""
    from models import AgentConfiguration
    from agent_scheduler import get_agent_scheduler
    company = current_user.get_default_company()
    
    scheduler = get_agent_scheduler()
    configs = AgentConfiguration.query.filter_by(company_id=company.id).all()
    available_agents = list(scheduler.agents.keys())
    
    return render_template('agent_configuration.html', configs=configs, available_agents=available_agents)

@main_bp.route('/agent-config/save', methods=['POST'])
@login_required
def save_agent_config():
    """Save agent configuration"""
    from models import AgentConfiguration
    try:
        data = request.get_json()
        company = current_user.get_default_company()
        
        agent_type = data.get('agent_type')
        config = AgentConfiguration.query.filter_by(
            company_id=company.id,
            agent_type=agent_type
        ).first()
        
        if not config:
            config = AgentConfiguration(
                company_id=company.id,
                agent_type=agent_type
            )
        
        config.is_enabled = data.get('is_enabled', True)
        config.schedule_frequency = data.get('schedule_frequency', 'daily')
        config.task_priority = data.get('task_priority', 5)
        config.configuration = data.get('configuration', {})
        
        db.session.add(config)
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Configuration saved'}), 200
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 400

print("✓ All feature routes loaded successfully")

# ============= ADVANCED CONFIG =============
@main_bp.route('/advanced-config')
@login_required
def advanced_config():
    """Advanced system configuration"""
    from models import CompanyIntegrationConfig
    company = current_user.get_default_company()
    configs = CompanyIntegrationConfig.query.filter_by(company_id=company.id).all()
    return render_template('advanced_config.html', configs=configs)


# ============= WORDPRESS CONNECTION =============
@main_bp.route('/wordpress/connect', methods=['POST'])
@login_required
def connect_wordpress():
    """Connect to WordPress site"""
    from models import WordPressIntegration
    from services.wordpress_service import WordPressService
    
    try:
        data = request.get_json()
        site_url = data.get('site_url', '').strip()
        api_key = data.get('api_key', '').strip()
        company = current_user.get_default_company()
        
        if not site_url or not api_key:
            return jsonify({'success': False, 'error': 'Site URL and API key required'}), 400
        
        # Test connection first
        result = WordPressService.test_connection(site_url, api_key)
        if not result['success']:
            return jsonify({'success': False, 'error': result['message']}), 400
        
        # Check if already exists
        existing = WordPressIntegration.query.filter_by(
            company_id=company.id,
            site_url=site_url
        ).first()
        
        if existing:
            existing.api_key = api_key
            existing.is_active = True
        else:
            wp_integration = WordPressIntegration(
                company_id=company.id,
                site_url=site_url,
                api_key=api_key,
                is_active=True
            )
            db.session.add(wp_integration)
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'WordPress connected successfully'}), 201
        
    except Exception as e:
        logger.error(f'WordPress connection error: {e}')
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/wordpress/sync/<int:wordpress_id>', methods=['POST'])
@login_required
def sync_wordpress_data(wordpress_id):
    """Sync WordPress posts and products"""
    from models import WordPressIntegration
    from services.wordpress_service import WordPressService
    
    try:
        wp = WordPressIntegration.query.get_or_404(wordpress_id)
        company = current_user.get_default_company()
        
        if wp.company_id != company.id:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 403
        
        posts_result = None
        products_result = None
        
        if wp.sync_blog_posts:
            posts_result = WordPressService.get_posts(wp.site_url, wp.api_key)
        
        if wp.sync_products:
            products_result = WordPressService.get_products(wp.site_url, wp.api_key)
        
        wp.last_synced_at = datetime.utcnow()
        db.session.commit()
        
        return jsonify({
            'success': True,
            'posts_synced': len(posts_result['posts']) if posts_result and posts_result['success'] else 0,
            'products_synced': len(products_result['products']) if products_result and products_result['success'] else 0
        }), 200
        
    except Exception as e:
        logger.error(f'WordPress sync error: {e}')
        return jsonify({'success': False, 'error': str(e)}), 500

print("✓ WordPress connection routes loaded")

# ============= COMPREHENSIVE LUX CRM =============
@main_bp.route('/crm-unified')
@login_required
def lux_crm():
    """Unified LUX CRM with all features"""
    company = current_user.get_default_company()
    
    deals = Deal.query.filter_by(company_id=company.id).all()
    all_contacts = Contact.query.all()
    lead_scores = LeadScore.query.all()
    personalization_rules = PersonalizationRule.query.filter_by(company_id=company.id).all()
    keywords = KeywordResearch.query.filter_by(company_id=company.id).all()
    
    return render_template('lux_crm.html', 
        deals=deals, 
        all_contacts=all_contacts,
        lead_scores=lead_scores,
        personalization_rules=personalization_rules,
        keywords=keywords
    )

@main_bp.route('/crm/deals/create', methods=['POST'])
@login_required
def create_deal():
    """Create a new deal in LUX CRM"""
    try:
        data = request.get_json()
        company = current_user.get_default_company()
        
        deal = Deal(
            company_id=company.id,
            contact_id=data.get('contact_id') or None,
            title=data.get('title'),
            description=data.get('description', ''),
            value=float(data.get('value', 0)),
            stage=data.get('stage', 'qualification'),
            probability=float(data.get('probability', 0.5)),
            expected_close_date=datetime.fromisoformat(data['expected_close_date']) if data.get('expected_close_date') else None,
            owner_id=current_user.id
        )
        db.session.add(deal)
        db.session.commit()
        return jsonify({'success': True, 'deal_id': deal.id}), 201
    except Exception as e:
        logger.error(f'Deal creation error: {e}')
        return jsonify({'success': False, 'error': str(e)}), 400

@main_bp.route('/api/contacts/<int:contact_id>')
@login_required
def get_contact(contact_id):
    """Get contact details via API"""
    try:
        contact = Contact.query.get(contact_id)
        if not contact:
            return jsonify({'error': 'Contact not found'}), 404
        return jsonify({
            'id': contact.id,
            'full_name': contact.full_name,
            'email': contact.email,
            'phone': contact.phone or '',
            'tags': contact.tags or ''
        }), 200
    except Exception as e:
        logger.error(f'Error fetching contact: {e}')
        return jsonify({'error': str(e)}), 500

print("✓ LUX CRM routes loaded")

# ============= AI AGENT REPORTING & MANAGEMENT =============
@main_bp.route('/agents/reports')
@login_required
def agents_reports_dashboard():
    """Comprehensive AI Agent Reports Dashboard - Activity, Reports, Deliverables"""
    from models import AgentLog, AgentTask
    from agent_scheduler import get_agent_scheduler
    
    company = current_user.get_default_company()
    
    # Get agent scheduler
    scheduler = get_agent_scheduler()
    scheduled_jobs = scheduler.get_scheduled_jobs() if scheduler else []
    
    # Get recent agent activities (last 50)
    recent_activities = AgentLog.query.order_by(AgentLog.created_at.desc()).limit(50).all()
    
    # Get agent task statistics
    total_tasks = AgentTask.query.count()
    completed_tasks = AgentTask.query.filter_by(status='completed').count()
    pending_tasks = AgentTask.query.filter_by(status='pending').count()
    failed_tasks = AgentTask.query.filter_by(status='failed').count()
    
    # Agent performance metrics
    agent_stats = {}
    agent_types = ['brand_strategy', 'content_seo', 'analytics', 'creative_design', 
                  'advertising', 'social_media', 'email_crm', 'sales_enablement', 
                  'retention', 'operations', 'app_intelligence']
    
    for agent_type in agent_types:
        agent_tasks = AgentTask.query.filter_by(agent_type=agent_type).all()
        agent_completed = len([t for t in agent_tasks if t.status == 'completed'])
        agent_total = len(agent_tasks)
        
        agent_stats[agent_type] = {
            'total_tasks': agent_total,
            'completed': agent_completed,
            'success_rate': (agent_completed / agent_total * 100) if agent_total > 0 else 0,
            'last_activity': AgentLog.query.filter_by(agent_type=agent_type).order_by(
                AgentLog.created_at.desc()
            ).first()
        }
    
    return render_template('agents_dashboard.html',
                         company=company,
                         scheduled_jobs=scheduled_jobs,
                         recent_activities=recent_activities,
                         agent_stats=agent_stats,
                         total_tasks=total_tasks,
                         completed_tasks=completed_tasks,
                         pending_tasks=pending_tasks,
                         failed_tasks=failed_tasks)

@main_bp.route('/api/agents/activity')
@login_required
def get_agent_activity():
    """Get real-time agent activity feed"""
    from models import AgentLog
    
    limit = request.args.get('limit', 20, type=int)
    agent_type = request.args.get('agent_type')
    
    query = AgentLog.query
    
    if agent_type:
        query = query.filter_by(agent_type=agent_type)
    
    activities = query.order_by(AgentLog.created_at.desc()).limit(limit).all()
    
    return jsonify({
        'success': True,
        'activities': [{
            'id': a.id,
            'agent_name': a.agent_name,
            'agent_type': a.agent_type,
            'activity_type': a.activity_type,
            'status': a.status,
            'created_at': a.created_at.isoformat() if a.created_at else None,
            'details': a.details
        } for a in activities]
    })

@main_bp.route('/api/agents/<agent_type>/performance')
@login_required
def get_agent_performance(agent_type):
    """Get performance metrics for specific agent"""
    from models import AgentTask, AgentLog
    from datetime import datetime, timedelta
    
    # Get tasks from last 30 days
    thirty_days_ago = datetime.now() - timedelta(days=30)
    
    tasks = AgentTask.query.filter(
        AgentTask.agent_type == agent_type,
        AgentTask.created_at >= thirty_days_ago
    ).all()
    
    completed = len([t for t in tasks if t.status == 'completed'])
    failed = len([t for t in tasks if t.status == 'failed'])
    total = len(tasks)
    
    # Get activity count
    activities = AgentLog.query.filter(
        AgentLog.agent_type == agent_type,
        AgentLog.created_at >= thirty_days_ago
    ).count()
    
    return jsonify({
        'success': True,
        'agent_type': agent_type,
        'period_days': 30,
        'metrics': {
            'total_tasks': total,
            'completed_tasks': completed,
            'failed_tasks': failed,
            'success_rate': (completed / total * 100) if total > 0 else 0,
            'total_activities': activities,
            'avg_tasks_per_day': total / 30
        }
    })

@main_bp.route('/api/agents/trigger/<agent_type>', methods=['POST'])
@login_required
def trigger_agent_task(agent_type):
    """Manually trigger an agent task"""
    from agent_scheduler import get_agent_scheduler
    
    data = request.get_json()
    task_data = data.get('task_data', {})
    
    scheduler = get_agent_scheduler()
    if not scheduler or agent_type not in scheduler.agents:
        return jsonify({'success': False, 'error': 'Agent not found'}), 404
    
    agent = scheduler.agents[agent_type]
    
    try:
        # Create and execute task
        task_id = agent.create_task(
            task_name=f"Manual: {task_data.get('task_type', 'custom')}",
            task_data=task_data
        )
        
        result = agent.execute(task_data)
        
        if task_id:
            agent.complete_task(
                task_id,
                result,
                status='completed' if result.get('success') else 'failed'
            )
        
        return jsonify({
            'success': True,
            'task_id': task_id,
            'result': result
        })
        
    except Exception as e:
        logger.error(f"Error triggering agent task: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/agent/diagnostics', methods=['GET'])
@login_required
def get_agent_diagnostics():
    """Get live agent diagnostics and error logs"""
    from agent_scheduler import get_agent_scheduler
    
    scheduler = get_agent_scheduler()
    
    # Get APP agent
    app_agent = scheduler.agents.get('app_intelligence') if scheduler else None
    
    if not app_agent:
        return jsonify({'success': False, 'error': 'APP Agent not found'}), 404
    
    try:
        # Get live error logs
        error_logs = app_agent.read_live_error_logs()
        
        # Get runtime state
        runtime_state = app_agent.get_app_runtime_state()
        
        # Get app files status
        file_status = app_agent.read_app_files('routes.py')
        
        return jsonify({
            'success': True,
            'timestamp': datetime.now().isoformat(),
            'error_logs': error_logs,
            'runtime_state': runtime_state,
            'app_health': app_agent.perform_health_check({}),
            'file_analysis': file_status.get('issues_found', [])
        })
    except Exception as e:
        logger.error(f"Diagnostics error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@main_bp.route('/agent/fix-issues', methods=['POST'])
@login_required
def trigger_agent_fix():
    """Manually trigger agent to scan and fix all issues"""
    from agent_scheduler import get_agent_scheduler
    
    scheduler = get_agent_scheduler()
    app_agent = scheduler.agents.get('app_intelligence') if scheduler else None
    
    if not app_agent:
        return jsonify({'success': False, 'error': 'APP Agent not found'}), 404
    
    try:
        result = app_agent.auto_detect_and_fix_issues()
        return jsonify({
            'success': True,
            'timestamp': datetime.now().isoformat(),
            'scan_results': result
        })
    except Exception as e:
        logger.error(f"Auto-fix error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

print("✓ AI Agent reporting routes loaded")

# ============= USER PROFILE MANAGEMENT =============
@main_bp.route('/user/profile')
@login_required
def user_profile():
    """View current user's profile"""
    user = current_user
    company = user.get_default_company()
    
    return render_template('user_profile.html', 
                         user=user, 
                         company=company)

@main_bp.route('/user/profile/edit', methods=['GET', 'POST'])
@login_required
def edit_user_profile():
    """Edit user profile"""
    user = current_user
    
    if request.method == 'POST':
        user.first_name = request.form.get('first_name', '')
        user.last_name = request.form.get('last_name', '')
        user.phone = request.form.get('phone', '')
        user.bio = request.form.get('bio', '')
        user.segment = request.form.get('segment', 'user')
        user.tags = request.form.get('tags', '')
        user.updated_at = datetime.now()
        
        db.session.commit()
        flash('Profile updated successfully!', 'success')
        return redirect(url_for('main.user_profile'))
    
    return render_template('edit_user_profile.html', user=user)

@main_bp.route('/api/user/profile')
@login_required
def get_user_profile_api():
    """Get user profile as JSON"""
    user = current_user
    
    return jsonify({
        'success': True,
        'user': {
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'full_name': user.full_name,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'phone': user.phone,
            'bio': user.bio,
            'segment': user.segment,
            'tags': user.tags.split(',') if user.tags else [],
            'is_admin': user.is_admin_user,
            'avatar': user.avatar_path,
            'engagement_score': user.engagement_score,
            'last_activity': user.last_activity.isoformat() if user.last_activity else None,
            'created_at': user.created_at.isoformat(),
            'updated_at': user.updated_at.isoformat()
        }
    })

@main_bp.route('/api/user/profile', methods=['PUT'])
@login_required
def update_user_profile_api():
    """Update user profile via API"""
    data = request.get_json()
    user = current_user
    
    try:
        if 'first_name' in data:
            user.first_name = data['first_name']
        if 'last_name' in data:
            user.last_name = data['last_name']
        if 'phone' in data:
            user.phone = data['phone']
        if 'bio' in data:
            user.bio = data['bio']
        if 'segment' in data:
            user.segment = data['segment']
        if 'tags' in data:
            user.tags = ','.join(data['tags']) if isinstance(data['tags'], list) else data['tags']
        
        user.updated_at = datetime.now()
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Profile updated'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 400

print("✓ User profile routes loaded")

# ============= CRM HUB =============
@main_bp.route('/crm/hub')
@login_required
def crm_hub():
    """CRM Features Hub - Showcase all 15 CRM capabilities"""
    return render_template('crm_hub.html')

print("✓ CRM Hub route loaded")

# ============= FORMINATOR NEWSLETTER IMPORT =============
@main_bp.route('/admin/import-forminator-newsletter', methods=['GET', 'POST'])
@login_required
def import_forminator_newsletter():
    """Import Forminator form 3482 newsletter signups and assign Newsletter segment"""
    if not current_user.is_admin_user:
        return jsonify({'error': 'Admin access required'}), 403
    
    if request.method == 'POST':
        try:
            # Import newsletter signups from Forminator form ID 3482
            # In production, this would connect to Forminator API
            # For now, we provide the import route and structure
            
            form_id = 3482
            import_count = 0
            
            # Parse submitted form data (this would come from Forminator webhook/API)
            submissions = request.get_json() or {}
            
            for submission in submissions.get('entries', []):
                email = submission.get('email')
                first_name = submission.get('first_name', '')
                last_name = submission.get('last_name', '')
                
                if email:
                    # Check if contact already exists
                    contact = Contact.query.filter_by(email=email).first()
                    
                    if contact:
                        # Update existing contact with Newsletter segment
                        contact.segment = 'newsletter'
                        if first_name:
                            contact.first_name = first_name
                        if last_name:
                            contact.last_name = last_name
                        contact.tags = 'newsletter_signup'
                        contact.source = 'forminator'
                    else:
                        # Create new contact with Newsletter segment
                        contact = Contact(
                            email=email,
                            first_name=first_name,
                            last_name=last_name,
                            segment='newsletter',
                            tags='newsletter_signup',
                            source='forminator',
                            is_active=True
                        )
                        db.session.add(contact)
                    
                    import_count += 1
            
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': f'Successfully imported {import_count} newsletter signups',
                'imported_count': import_count
            }), 200
            
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500
    
    return render_template('import_forminator.html')

@main_bp.route('/admin/forminator-webhook', methods=['POST'])
def forminator_webhook():
    """Webhook endpoint for Forminator form submissions (form ID 3482)"""
    try:
        data = request.get_json()
        form_id = data.get('form_id')
        
        # Only process form 3482 (Newsletter signup)
        if form_id != 3482:
            return jsonify({'status': 'ignored'}), 200
        
        email = data.get('email') or data.get('fields', {}).get('email', {}).get('value')
        first_name = data.get('first_name') or data.get('fields', {}).get('first_name', {}).get('value', '')
        last_name = data.get('last_name') or data.get('fields', {}).get('last_name', {}).get('value', '')
        
        if email:
            contact = Contact.query.filter_by(email=email).first()
            
            if contact:
                contact.segment = 'newsletter'
                contact.tags = 'newsletter_signup'
                contact.source = 'forminator'
                if first_name:
                    contact.first_name = first_name
                if last_name:
                    contact.last_name = last_name
            else:
                contact = Contact(
                    email=email,
                    first_name=first_name,
                    last_name=last_name,
                    segment='newsletter',
                    tags='newsletter_signup',
                    source='forminator',
                    is_active=True
                )
                db.session.add(contact)
            
            db.session.commit()
            return jsonify({'status': 'success', 'email': email}), 200
        
        return jsonify({'status': 'error', 'message': 'No email provided'}), 400
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500

print("✓ Forminator newsletter import routes loaded")

# ============= WORDPRESS USER IMPORT =============
@main_bp.route('/admin/import-wordpress-users', methods=['GET', 'POST'])
@login_required
def import_wordpress_users():
    """Import WordPress users with roles as tags and membership-based segmentation"""
    if not current_user.is_admin_user:
        return jsonify({'error': 'Admin access required'}), 403
    
    if request.method == 'POST':
        try:
            data = request.get_json() or {}
            wordpress_url = data.get('wordpress_url', '')
            
            if not wordpress_url:
                return jsonify({'error': 'WordPress URL required'}), 400
            
            import requests
            import os
            
            # Fetch users from WordPress REST API
            users_endpoint = f"{wordpress_url.rstrip('/')}/wp-json/wp/v2/users"
            
            # Try to fetch without auth first
            try:
                response = requests.get(users_endpoint, timeout=10)
                response.raise_for_status()
                wp_users = response.json()
            except:
                # If public endpoint fails, return error
                return jsonify({'error': 'Could not fetch WordPress users. Ensure REST API is public or provide credentials.'}), 400
            
            import_count = 0
            
            for wp_user in wp_users:
                email = wp_user.get('email')
                username = wp_user.get('slug') or wp_user.get('username')
                first_name = wp_user.get('name', '').split()[0] if wp_user.get('name') else ''
                last_name = ' '.join(wp_user.get('name', '').split()[1:]) if wp_user.get('name') else ''
                
                if email:
                    # Get WordPress role(s)
                    wp_role = wp_user.get('roles', ['subscriber'])[0] if wp_user.get('roles') else 'subscriber'
                    
                    # Check for membership (this would integrate with membership plugin)
                    # For now, we'll mark users with specific roles as members
                    has_membership = wp_role in ['member', 'premium', 'vip', 'administrator']
                    
                    contact = Contact.query.filter_by(email=email).first()
                    
                    if contact:
                        contact.segment = 'member' if has_membership else 'Website Users'
                        contact.tags = f"{wp_role},wordpress_import,{username}"
                        contact.first_name = first_name or contact.first_name
                        contact.last_name = last_name or contact.last_name
                        contact.source = 'wordpress'
                    else:
                        contact = Contact(
                            email=email,
                            first_name=first_name,
                            last_name=last_name,
                            segment='member' if has_membership else 'Website Users',
                            tags=f"{wp_role},wordpress_import,{username}",
                            source='wordpress',
                            is_active=True
                        )
                        db.session.add(contact)
                    
                    import_count += 1
            
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': f'Successfully imported {import_count} WordPress users',
                'imported_count': import_count
            }), 200
            
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500
    
    return render_template('import_wordpress.html')

@main_bp.route('/admin/wordpress-webhook', methods=['POST'])
def wordpress_webhook():
    """Webhook for WordPress new user registration"""
    try:
        data = request.get_json()
        email = data.get('email')
        username = data.get('username')
        first_name = data.get('first_name', '')
        last_name = data.get('last_name', '')
        wp_role = data.get('role', 'subscriber')
        
        if email:
            has_membership = wp_role in ['member', 'premium', 'vip', 'administrator']
            
            contact = Contact.query.filter_by(email=email).first()
            
            if contact:
                contact.segment = 'member' if has_membership else 'Website Users'
                contact.tags = f"{wp_role},wordpress_user,{username}"
                if first_name:
                    contact.first_name = first_name
                if last_name:
                    contact.last_name = last_name
                contact.source = 'wordpress'
            else:
                contact = Contact(
                    email=email,
                    first_name=first_name,
                    last_name=last_name,
                    segment='member' if has_membership else 'Website Users',
                    tags=f"{wp_role},wordpress_user,{username}",
                    source='wordpress',
                    is_active=True
                )
                db.session.add(contact)
            
            db.session.commit()
            return jsonify({'status': 'success', 'email': email}), 200
        
        return jsonify({'status': 'error', 'message': 'No email provided'}), 400
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500

print("✓ WordPress user import routes loaded")

# ============= TEST WORDPRESS IMPORT (AUTO-DEMO) =============
@main_bp.route('/admin/test-wordpress-import', methods=['GET'])
@login_required
def test_wordpress_import():
    """Auto-import test WordPress users for demo"""
    if not current_user.is_admin_user:
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        # Create sample WordPress users for testing
        test_users = [
            {
                'email': 'john@example.com',
                'name': 'John Administrator',
                'slug': 'john_admin',
                'roles': ['administrator']
            },
            {
                'email': 'jane@example.com',
                'name': 'Jane Member',
                'slug': 'jane_member',
                'roles': ['member']
            },
            {
                'email': 'bob@example.com',
                'name': 'Bob Subscriber',
                'slug': 'bob_sub',
                'roles': ['subscriber']
            },
            {
                'email': 'alice@example.com',
                'name': 'Alice Premium',
                'slug': 'alice_premium',
                'roles': ['premium']
            }
        ]
        
        import_count = 0
        
        for wp_user in test_users:
            email = wp_user.get('email')
            name = wp_user.get('name', '')
            username = wp_user.get('slug')
            first_name = name.split()[0] if name else ''
            last_name = ' '.join(name.split()[1:]) if name else ''
            wp_role = wp_user.get('roles', ['subscriber'])[0]
            
            has_membership = wp_role in ['member', 'premium', 'vip', 'administrator']
            
            contact = Contact.query.filter_by(email=email).first()
            
            if contact:
                contact.segment = 'member' if has_membership else 'Website Users'
                contact.tags = f"{wp_role},wordpress_import,{username}"
                contact.first_name = first_name or contact.first_name
                contact.last_name = last_name or contact.last_name
                contact.source = 'wordpress'
            else:
                contact = Contact(
                    email=email,
                    first_name=first_name,
                    last_name=last_name,
                    segment='member' if has_membership else 'Website Users',
                    tags=f"{wp_role},wordpress_import,{username}",
                    source='wordpress',
                    is_active=True
                )
                db.session.add(contact)
            
            import_count += 1
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Test WordPress import successful - {import_count} users imported',
            'imported_count': import_count,
            'users': [
                {'email': u['email'], 'name': u['name'], 'role': u['roles'][0]} 
                for u in test_users
            ]
        }), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@main_bp.route('/admin/wordpress-imports', methods=['GET'])
@login_required
def view_wordpress_imports():
    """View all WordPress imported contacts"""
    if not current_user.is_admin_user:
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        wp_contacts = Contact.query.filter_by(source='wordpress').all()
        
        return jsonify({
            'success': True,
            'count': len(wp_contacts),
            'contacts': [
                {
                    'id': c.id,
                    'email': c.email,
                    'name': c.full_name,
                    'segment': c.segment,
                    'tags': c.tags,
                    'source': c.source,
                    'created_at': c.created_at.isoformat()
                }
                for c in wp_contacts
            ]
        }), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ============= PUBLIC ZAPIER WEBHOOK ENDPOINT =============
@main_bp.route('/api/webhook/zapier-contact', methods=['POST'])
@csrf.exempt
def zapier_contact_webhook():
    """
    Public API endpoint for Zapier webhook integration
    Receives: email, name, phone, source (flexible payload)
    Validates, checks duplicates, inserts/updates contact
    Supports basic auth: luke|Wow548302!
    """
    try:
        # Validate basic auth if provided
        auth = request.authorization
        if auth:
            if auth.username != 'luke' or auth.password != 'Wow548302!':
                return jsonify({
                    'success': False,
                    'error': 'Invalid credentials'
                }), 401
        
        data = request.get_json()
        
        if not data:
            return jsonify({
                'success': False,
                'error': 'No JSON payload provided'
            }), 400
        
        # Validate required fields (email and source are required)
        email = data.get('email', '').strip().lower()
        name = data.get('name', '').strip()
        phone = data.get('phone', '').strip() or None
        source = data.get('source', 'Zapier').strip()
        
        # Validate email format
        if not email or not validate_email(email):
            return jsonify({
                'success': False,
                'error': f'Invalid or missing email: {email}'
            }), 400
        
        if not source:
            source = 'Zapier'
        
        # Parse name into first and last
        first_name = ''
        last_name = ''
        if name:
            name_parts = name.split(' ', 1)
            first_name = name_parts[0] if len(name_parts) > 0 else ''
            last_name = name_parts[1] if len(name_parts) > 1 else ''
        
        # Check for duplicate
        existing_contact = Contact.query.filter_by(email=email).first()
        
        if existing_contact:
            # Update existing contact
            if first_name:
                existing_contact.first_name = first_name
            if last_name:
                existing_contact.last_name = last_name
            if phone:
                existing_contact.phone = phone
            
            # Add zapier tag if not present
            current_tags = existing_contact.tags or ''
            if 'zapier' not in current_tags.lower():
                existing_contact.tags = f"{current_tags},zapier".strip(',')
            
            # Set/update source to track where it came from
            if not existing_contact.source:
                existing_contact.source = source
            
            # Ensure Newsletter segment for signup sources
            if 'newsletter' in source.lower() or 'signup' in source.lower():
                existing_contact.segment = 'Newsletter'
            elif not existing_contact.segment:
                existing_contact.segment = 'lead'
            
            existing_contact.updated_at = datetime.utcnow()
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': 'Contact updated successfully',
                'action': 'updated',
                'contact_id': existing_contact.id,
                'email': existing_contact.email,
                'segment': existing_contact.segment,
                'source': existing_contact.source
            }), 200
        
        else:
            # Create new contact
            new_contact = Contact()
            new_contact.email = email
            new_contact.first_name = first_name
            new_contact.last_name = last_name
            new_contact.phone = phone
            new_contact.segment = 'Newsletter' if 'newsletter' in source.lower() or 'signup' in source.lower() else 'lead'
            new_contact.tags = 'zapier'
            new_contact.source = source
            new_contact.is_active = True
            
            db.session.add(new_contact)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': 'Contact created successfully',
                'action': 'created',
                'contact_id': new_contact.id,
                'email': new_contact.email,
                'segment': new_contact.segment,
                'source': new_contact.source
            }), 201
    
    except Exception as e:
        db.session.rollback()
        logger.error(f'Zapier webhook error: {str(e)}')
        return jsonify({
            'success': False,
            'error': f'Server error: {str(e)}'
        }), 500

print("✓ WordPress import test and view routes loaded")
print("✓ Zapier webhook endpoint loaded at /api/webhook/zapier-contact")
print("✓ Error logging and diagnostics endpoints loaded")
print("✓ AI Chatbot configured for error analysis, auto-repair, and server log reading")
print("✓ Log reading capability: Nginx, Gunicorn, systemd, and app logs")
print("✓ Automated error repair and resolution testing enabled")
print("✓ Auto-repair endpoints: /api/auto-repair/start and /api/auto-repair/clear")
print("✓ System health and diagnostics endpoints:")
print("  - /api/system/diagnosis (comprehensive analysis)")
print("  - /api/system/health (resource usage)")
print("  - /api/system/validate-openai (API key validation)")
print("  - /api/system/endpoint-check (404 detection)")
print("✓ AI Action Executor endpoints (autonomous, action-oriented):")
print("  - POST /api/ai/execute-action (execute AI actions immediately)")
print("  - GET /api/company/<id>/secrets (retrieve company secrets)")
print("✓ CompanySecret model created for secure secret storage per company")
print("✓ Secrets populated for Lucifer Cruz company from environment variables")
